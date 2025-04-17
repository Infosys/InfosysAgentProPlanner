''' 
Copyright 2025 Infosys Ltd.
 
Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated documentation files (the "Software"), 
to deal in the Software without restriction, including without limitation the rights to use, copy, modify, merge, publish, distribute, sublicense, 
and/or sell copies of the Software, and to permit persons to whom the Software is furnished to do so, subject to the following conditions:
 
The above copyright notice and this permission notice shall be included in all copies 
or substantial portions of the Software.
 
THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, 
INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE 
AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, 
DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM, 
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.
'''


import requests
from openai import AzureOpenAI
from typing import Dict, List, Union, Any, Optional
from fastapi import HTTPException
from pydantic import BaseModel, field_validator
import importlib
import os
import sys
from datetime import datetime
import json
import pickle
import numpy as np
import google.generativeai as genai
from rank_bm25 import BM25Okapi
import faiss
from sklearn.feature_extraction.text import TfidfVectorizer
import re
import logging
from enum import Enum
from docx import Document
from io import BytesIO
import uuid
import asyncio

# --- Data Models (Pydantic) ---
os.chdir(os.path.dirname(os.path.abspath(__file__)))

INSTRUCTIONS_FILE = "Instructions.docx"  # Update if needed
EMBEDDINGS_FILE = "instruction_embeddings.pkl"  # Update if needed
DYNAMIC_EMBEDDINGS_FILE = "instruction_embeddings.pkl"

class SafeUnpickler(pickle.Unpickler):
    """
    A restricted unpickler that only allows safe built-in types and specific external classes to be deserialized.
    """
    def find_class(self, module, name):
        # Allow only safe built-in types and specific external classes
        allowed_classes = {
            "builtins": {"dict", "list", "str", "int", "float", "bool"},
            "rank_bm25": {"BM25Okapi"},  # Allow rank_bm25.BM25Okapi
            "faiss.swigfaiss_avx2": {"IndexFlatIP"}  # Allow faiss.swigfaiss_avx2.IndexFlatIP
        }
        if module in allowed_classes and name in allowed_classes[module]:
            return super().find_class(module, name)
        raise pickle.UnpicklingError(f"Attempted to load unsafe class: {module}.{name}")

def safe_load_pickle(file_path: str) -> Dict[str, Any]:
    """
    Safely loads a pickle file using a restricted unpickler.
    """
    try:
        with open(file_path, "rb") as f:
            return SafeUnpickler(f).load()
    except pickle.UnpicklingError as e:
        raise ValueError(f"Unsafe deserialization attempt detected: {e}")
    except Exception as e:
        raise ValueError(f"Error loading pickle file: {e}")

class ToolExecutionResult(BaseModel):
    """Represents the result of a tool execution."""
    output: Optional[Any]
    error: Optional[str]

class StepExecutionResult(BaseModel):
    """Represents the result of a single step execution."""
    tool_name: str
    arguments: Dict[str, Any]
    result: ToolExecutionResult

    @field_validator("arguments", mode="before")  # Use mode="before"
    def make_arguments_serializable(cls, v):
        return {k: v[k] for k in v if k not in ["memory_manager", "llm_manager"]}


class PlanExecutionResult(BaseModel):
    """Represents the result of the entire plan execution."""
    steps: List[StepExecutionResult]
    memory: Dict[str, Any]

class GeneratePlanRequest(BaseModel):
    """Request body for generating a plan."""
    goal: str
    plan_verifier: Optional[bool] = False

class ExecutePlanRequest(BaseModel):
    """Request body for executing a plan."""
    plan: List[str]
    arguments_verifier: Optional[bool] = False

class RunAgentRequest(BaseModel):
    """Request body for running the agent."""
    goal: Optional[str] = None
    tool_verifier: Optional[bool] = False
    plan_verifier: Optional[bool] = False
    arguments_verifier: Optional[bool] = False
    unstruct_input: Optional[str] = None # Add a field for user input
    goals: Optional[List[str]] = None
    auto_verifier: Optional[bool] = False  # Add the auto_verifier flag
    split_execution: Optional[bool] = False
    auto_learning: Optional[bool] = False  # Add the auto_verifier flag
    llm_type: Optional[str] = None  # Add llm_type field here
    agents: Optional[List[str]] = None
    tool_result_verifier: Optional[bool] = False # NEW FIELD
    result_verifier: Optional[bool] = False # NEW FIELD
    human_tool_result_verifier: Optional[bool] = False
    

class MessageType(Enum):
    PROMPT = "prompt"
    MESSAGE = "message"
    RESULT = "result"
    PLAN = "plan"
    LOG = "log"  # Add a new message type for logs


# --- Memory Manager ---
class MemoryManager:
    def __init__(self):
        self.memory = {}  # Long-term memory
        self.results = {} # Short-term, per-plan execution results
        self.conversation_history = []

    def store_result(self, node_id: str, result: Dict):
        """Stores the result of a tool execution."""
        self.results[node_id] = result

    def get_result(self, node_id: str) -> Any:
        """Retrieves the result of a previous tool execution."""
        return self.results.get(node_id)

    def store_in_memory(self, key: str, value: Any):
        """Stores a value in the long-term memory."""
        self.memory[key] = value

    def get_from_memory(self, key: str) -> Any:
        """Retrieves a value from the long-term memory."""
        return self.memory.get(key)

    def clear_results(self):
        """Clears the short-term results."""
        self.results = {}

    def clear_memory(self):
        """Clears the long-term memory."""
        self.memory = {}
    
    def add_to_conversation(self, entry: Dict):
        """Adds an entry to the conversation history.

        Args:
            entry: A dictionary representing a single turn in the conversation.
                   Should include keys like: 'user_input', 'inferred_goal',
                   'plan', 'results', 'feedback'.  Not all keys are required
                   for every entry.
        """
        self.conversation_history.append(entry)

    def get_conversation_history(self) -> List[Dict]:
        """Returns the entire conversation history."""
        return self.conversation_history

    def clear_conversation_history(self):
        """Clears the conversation history."""
        self.conversation_history = []


    def _substitute_placeholders(self, arguments: Dict) -> Dict:
      """
      Substitutes placeholders in the argument values with results from previous steps.

      Args:
          arguments: The arguments dictionary.

      Returns:
          A new dictionary with placeholders replaced by actual values.
      """
      substituted_arguments = {}
      for arg_name, arg_value in arguments.items():
          if isinstance(arg_value, str):
              # Iterate through the ENTIRE results dictionary.
              for key, value in self.results.items():
                  placeholder = f"{{{key}}}"
                  if placeholder in arg_value:  # Check if the placeholder exists
                      if isinstance(value,dict) and "result" in value:
                          arg_value = arg_value.replace(placeholder, str(value["result"]))
                      elif isinstance(value, (str, int, float, list)):  # Handle direct values
                          arg_value = arg_value.replace(placeholder, str(value))

                      else:
                          logging.warning(f"Skipping placeholder substitution for key '{key}' in argument '{arg_name}'. Result value is not a dictionary with 'result' key (or is missing). Value type: {type(value)}") # Log warning if value is not as expected
              substituted_arguments[arg_name] = arg_value
          elif isinstance(arg_value, list):
                substituted_list = []
                for item in arg_value:
                    if isinstance(item, str):
                        for key, value in self.results.items():
                            placeholder = f"{{{key}}}"
                            if placeholder in item:
                                if isinstance(value, dict) and "result" in value:
                                  item = item.replace(placeholder, str(value["result"]))
                                elif isinstance(value, (str, int, float, list)): # Handle lists
                                  item = item.replace(placeholder, str(value))
                                else:
                                  logging.warning(f"Skipping placeholder substitution for '{placeholder}' in list.  Value type not supported: {type(value)}")
                        substituted_list.append(item)
                    else:
                        substituted_list.append(item) #if not of type string, keep original value
                substituted_arguments[arg_name] = substituted_list
          else:
              substituted_arguments[arg_name] = arg_value

      return substituted_arguments

    async def _llm_substitute_placeholders(self, arguments: Dict, llm_manager) -> Dict:
        """
        Substitutes placeholders in arguments using an LLM for potentially smarter substitution.

        Args:
            arguments: The arguments dictionary (can contain placeholders).
            results: The results dictionary from previous steps.
            llm_manager: The LLMManager instance.

        Returns:
            A new dictionary with placeholders substituted by LLM.
        """
        substituted_arguments = {}
        for arg_name, arg_value in arguments.items():
            if isinstance(arg_value, str) and "{{" in arg_value and "}}" in arg_value: # Check if it's a string and contains placeholders
                prompt = f"""
                You are an AI assistant specialized in substituting placeholders in argument values.
                Your task is to replace placeholders in the format `{{{{node_name}}}}` within a given argument value string with the corresponding results from a `results` dictionary.

                **Argument Name:** {arg_name}
                **Argument Value (with placeholders):** {arg_value}
                **Results from Previous Steps (JSON):**
                {json.dumps(self.results)}

                **Instructions:**
                1.  Identify all placeholders in the `Argument Value`. Placeholders are enclosed in double curly braces, like `{{{{node_name}}}}`.
                2.  For each placeholder `{{{{node_name}}}}`, find the corresponding `node_name` in the `Results from Previous Steps`.
                3.  If a `node_name` is found in the `Results`, extract the `result` associated with that `node_name` from the `Results`.
                4.  Replace the placeholder `{{{{node_name}}}}` in the `Argument Value` with the extracted `result` (converted to a string).
                5.  If a placeholder `{{{{node_name}}}}` refers to a `node_name` that is **NOT** found in the `Results from Previous Steps`, **leave the placeholder as is** in the output and include a warning message in your response (but still return valid JSON). Do not remove the placeholder if not found, return as is.
                6.  Maintain the original string format and any surrounding text of the `Argument Value`, only replace the placeholders.
                7.  Return the **modified Argument Value as a JSON string**.  If no substitution is needed, return the original Argument Value as a JSON string.

                **Example:**

                **Argument Value (with placeholders):** "The predicted payment day is {{{{predict_payment_timing_1}}}} for customer {{{{get_customer_code_1}}}}"
                **Results from Previous Steps (JSON):**
                {{
                    "predict_payment_timing_1": {{"node_id": "predict_payment_timing_1", "result": "Wednesday"}},
                    "get_customer_code_1": {{"node_id": "get_customer_code_1", "result": "ACME_CODE"}}
                }}

                **Output (JSON string):**
                "The predicted payment day is Wednesday for customer ACME_CODE"

                **Now, process the Argument Value provided and return the substituted value as a JSON string:**
                """

                llm_response_str = llm_manager.llm_call(prompt, "gemini") # Assuming default to gemini for reasoning
                try:
                    # Directly use the string response as the substituted value (no need to parse JSON in response)
                    substituted_value = llm_response_str.strip()

                except Exception as e:
                    logging.warning(f"LLM-based substitution failed for argument '{arg_name}'. Using original value. Error: {e}")
                    substituted_value = arg_value # Fallback to original value on error
            else:
                substituted_value = arg_value # No substitution needed, use original value

            substituted_arguments[arg_name] = substituted_value

        return substituted_arguments



def load_config(config_file: str = "config.json") -> Dict:
    """Loads configuration from a JSON file."""
    try:
        with open(config_file, "r") as f:
            config = json.load(f)
        return config
    except (FileNotFoundError, KeyError) as e:
        logging.error(f"Error loading configuration: {e}")
        raise ValueError(f"Error loading configuration: {e}")

# --- LLM Manager ---
#model names: deepseek-r1:14b deepseek-r1:1.5b
class LLMManager:
    def __init__(self, agent,api_key, endpoint, model_deployment_name, gemini_api_key, gemini_model_name, deepseek_api_key=None, deepseek_model='deepseek-r1',deepseek_model_local='deepseek-r1:8b', default_llm="gemini"):
        self.agent = agent
        self.default_llm = default_llm
        self.gemini_api_key = gemini_api_key
        self.reasoning_model = 'gemini-2.0-flash-thinking-exp'  # Add reasoning model
        self.reasoning_config = {'include_thoughts': True}  # Special config
        self.deepseek_model_local = deepseek_model_local
       

        #if self.default_llm == "azure_openai":
        self.azure_openai_client = AzureOpenAI(
            api_key=api_key,
            api_version="2024-02-01",
            azure_endpoint=endpoint
        )
        self.model_deployment_name = model_deployment_name
        #elif self.default_llm == "gemini":
        genai.configure(api_key=gemini_api_key)
        self.gemini_model = genai.GenerativeModel(gemini_model_name)

        # New Deepseek initialization
        """self.deepseek_client = None
        self.deepseek_model = deepseek_model
        if deepseek_api_key:
            from deepseek import DeepseekClient 
            self.deepseek_client = DeepseekClient(api_key=deepseek_api_key)"""""

        #else:
            #raise ValueError(f"Invalid default_llm: {default_llm}. Must be 'azure_openai' or 'gemini'.")
    
    
    def get_client_and_model(self, llm_type: str = None):
        """Returns the appropriate client and model based on llm_type."""
        llm_to_use = llm_type if llm_type is not None else self.default_llm
        
        if llm_to_use == "azure_openai":
            return self.azure_openai_client, self.model_deployment_name
        elif llm_to_use == "gemini":
            return self.gemini_model, None  # No separate model name for Gemini
        elif llm_to_use == "gemini_reasoning":  # New condition
            return self.gemini_model, 'gemini-2.0-flash-thinking-exp'
        #elif llm_to_use == "deepseek_r1":  # New condition
            #return self.deepseek_client, self.deepseek_model
        elif llm_to_use == "deepseek_r1":
            return None, self.deepseek_model_local
        else:
            raise ValueError(f"Invalid llm_type: {llm_to_use}. Must be 'azure_openai' or 'gemini'.")

    def llm_call(self, prompt: str, llm_type: str) -> Union[Dict, List, str, None]:
        """Common utility to call an LLM and handle JSON parsing."""
        client, model_name = self.get_client_and_model(llm_type)

        if llm_type == "deepseek_r1":
            try:
                return self._query_deepseek_local(prompt, model_name)
            except Exception as e:
                error_message = f"Deepseek local error: {e}" # Create error message
                logging.error(error_message)
                # asyncio.create_task(self._send_ui_error_message(error_message)) # Use asyncio.create_task to avoid blocking llm_call
                return None

        if llm_type == "gemini_reasoning_old":
            try:
                response = genai.generate_content(
                    model=self.reasoning_model,
                    contents=[{'parts': [{'text': prompt}]}],
                    generation_config=genai.GenerationConfig(**self.reasoning_config)
                )
                return response.text.strip()
            except Exception as e:
                error_message = f"Gemini Reasoning error (old): {e}" # Create error message
                logging.error(error_message)
                # asyncio.create_task(self._send_ui_error_message(error_message))
                return None
            
        if llm_type == "gemini_reasoning":
            try:
                model = genai.GenerativeModel(self.reasoning_model)
                response = model.generate_content(prompt)
                return response.text.strip()
            except Exception as e:
                error_message = f"Gemini Reasoning error: {e}" # Create error message
                logging.error(error_message)
                # asyncio.create_task(self._send_ui_error_message(error_message))
                return None

        if llm_type == "gemini":
            messages = [{"role": "user", "parts": [prompt]}]
            try:
                response = client.generate_content(messages)
                response_str = response.text.strip()
                logging.info(f"Gemini Raw response: {response_str}")
            except Exception as e:
                error_message = f"Error calling Gemini LLM: {e}" # Create error message
                logging.error(error_message)
                # asyncio.create_task(self._send_ui_error_message(error_message))
                return None
        else:  # Azure OpenAI
            messages = [
                {"role": "system", "content": "You are an AI assistant."},
                {"role": "user", "content": prompt}
            ]
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=messages,
                    temperature=0.2
                )
                response_str = response.choices[0].message.content.strip()
            except Exception as e:
                error_message = f"Error calling Azure OpenAI: {e}" # Create error message
                logging.error(error_message)
                # asyncio.create_task(self._send_ui_error_message(error_message))
                return None

        # Cleanup JSON output
        try: # Wrap in try-except block
            if response_str.startswith("```json"):
                response_str = response_str[7:]
            if response_str.endswith("```"):
                response_str = response_str[:-3]
        except Exception as e:
            error_message = f"Error during JSON cleanup in llm_call: {e}"
            logging.error(error_message)
            # asyncio.create_task(self._send_ui_error_message(error_message))
            return None # Return None in case of error during cleanup.

        return response_str
    
    
    def _query_deepseek_local(self, prompt: str, model: str) -> str:
        """Direct API call to local Ollama instance"""
        url = "http://127.0.0.1:11434/api/generate"
        payload = {
            "prompt": prompt,
            "model": model,
            "stream": False,
            "options": {
                "temperature": 0.3,
                "num_predict": 2000
            }
        }
        
        response = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, stream=True)
        
        if response.status_code != 200:
            raise ConnectionError(f"API request failed: {response.status_code} - {response.text}")

        full_response = []
        for line in response.iter_lines():
            if line:
                try:
                    json_line = json.loads(line.decode('utf-8'))
                    response_text = json_line.get("response", "")
                    # Remove <think> tags and their content
                    response_text = re.sub(r'<think>.*?</think>', '', response_text, flags=re.DOTALL)
                    full_response.append(response_text)
                except json.JSONDecodeError as e:
                    logging.warning(f"Failed to decode line: {e}")
        
        result = ''.join(full_response).strip()

        # Cleanup JSON output
        if result.startswith("```json"):
            result = result[7:]
        if result.endswith("```"):
            result = result[:-3]

        return result
                    
        #return ''.join(full_response)

    def infer_goal_from_text(self, text: str, llm_type: str, tools: Dict, instruction_retriever: Any, split_execution: bool = False, selected_agents: List[str] = None, previous_context: str = None) -> List[str]:
        """
        Infers the user's goal from unstructured text using an LLM.
        """
        client, model_name = self.get_client_and_model(llm_type)
        tool_descriptions = _get_tool_descriptions(tools, selected_agents=selected_agents)  # Get tool descriptions
        relevant_instructions = instruction_retriever.get_relevant_instructions(text, num_instructions=3, similarity_threshold=0.7, api_key=self.gemini_api_key)

        # Format instructions for the prompt, handling missing fields
        instructions_str = ""
        if relevant_instructions:
            for instruction in relevant_instructions:
                instructions_str += "**Instruction:**\n"
                if "Input" in instruction:
                    instructions_str += f"  Input: {instruction['Input']}\n"
                if "Goal" in instruction:
                    instructions_str += f"  Goal: {instruction['Goal']}\n"
                if "Tools to be used" in instruction:
                    instructions_str += f"  Tools to be used: {', '.join(instruction['Tools'])}\n"
                if "Sequence of Tools" in instruction:
                    instructions_str += f"  Sequence of Tools: {', '.join(instruction['Sequence'])}\n"
                if instruction.get("Additional Information"):
                    instructions_str += f"  Additional Information: {instruction['Additional Information']}\n"
                instructions_str += "---\n"

        print("entering infer goal: ",text)

        prompt = f"""You are an AI assistant that helps extract the user goal from unstructured text inputs.

    **Input:**
    {text}

    **Available Tools and their Descriptions:**
        {tool_descriptions}

    **Previous Conversation Context (if any):**
    {previous_context if previous_context else "No prior conversation."}  # Include context in prompt

    **Instructions:**
    - Analyze the input text and identify the user's underlying goal.
    - Ensure to maintain the values or arguments given in the input as part of the goal
    - The goal should describe the desired outcome or action to be performed.
    - Focus on the primary objective(s).
    - **Only generate goal that can be achieved using the available tools listed above.**

    **Relevant Instructions based on your input:**
    {instructions_str}

        **Examples:**

        Input: "I received an email from customer ABNMST saying that his payment is due next week. Can you predict his payment date and draft a reminder email?"
        Available Tools: predict_payment_timing, draft_email, summarize_text, multiply_numbers, divide_numbers, calculate_sum, convert_to_bullets
        Output: ["Predict payment date for customer ABNMST and Draft a reminder email for customer ABNMST"]

        Input: "My computer is running very slow and I keep getting a blue screen error. Can you help?"
        Available Tools: troubleshoot_computer, run_diagnostics, restart_system
        Output: ["Troubleshoot slow computer performance and check the blue screen error"]

        Input: "Summarize document.txt and translate it to Spanish"
        Available Tools: summarize_text, translate_text, extract_text
        Output: ["Summarize document.txt and Translate the summary to Spanish"]

        **Output (JSON array of strings):**
    """
        arguments_str = self.llm_call(prompt, llm_type)

        if not arguments_str:
            logging.error("LLM returned an empty response.")
            return []

        
        inferred_goals = json.loads(arguments_str)
        return inferred_goals

    def infer_goal_from_text_Split(self, text: str, llm_type: str, tools: Dict, instruction_retriever: Any, split_execution: bool = False, selected_agents: List[str] = None,previous_context: str = None) -> List[str]:
        """
        Infers the user's goal from unstructured text using an LLM.
        """
        client, model_name = self.get_client_and_model(llm_type)
        tool_descriptions = _get_tool_descriptions(tools, selected_agents=selected_agents)  # Get tool descriptions
        relevant_instructions = instruction_retriever.get_relevant_instructions(text, num_instructions=3, similarity_threshold=0.7, api_key=self.gemini_api_key)

        # Format instructions for the prompt, handling missing fields
        instructions_str = ""
        if relevant_instructions:
            for instruction in relevant_instructions:
                instructions_str += "**Instruction:**\n"
                if "Input" in instruction:
                    instructions_str += f"  Input: {instruction['Input']}\n"
                if "Goal" in instruction:
                    instructions_str += f"  Goal: {instruction['Goal']}\n"
                if "Tools to be used" in instruction:
                    instructions_str += f"  Tools to be used: {', '.join(instruction['Tools'])}\n"
                if "Sequence of Tools" in instruction:
                    instructions_str += f"  Sequence of Tools: {', '.join(instruction['Sequence'])}\n"
                if instruction.get("Additional Information"):
                    instructions_str += f"  Additional Information: {instruction['Additional Information']}\n"
                instructions_str += "---\n"

        print("entering infer goal: ",text)

        prompt = f"""You are an AI assistant that helps extract user goals from unstructured text inputs.

    **Input:**
    {text}

    **Available Tools and their Descriptions:**
        {tool_descriptions}

     **Previous Conversation Context (if any):**
        {previous_context if previous_context else "No prior conversation."}  

    **Instructions:**
    - Analyze the input text and identify the user's underlying goal(s).
    - Ensure to maintain the values or arguments given in the input as part of the goal or goals.
    - The goal(s) should describe the desired outcome or action to be performed.
    - Express each goal in a clear and concise sentence.
    - Focus on the primary objective(s).
    - **Only generate goals that can be achieved using the available tools listed above.**
    - **If the input implies multiple distinct goals, generate a separate goal for each one.**
    - **Return the goals as a JSON array of strings, where each string is a separate goal.**

    **Relevant Instructions based on your input:**
    {instructions_str}

        **Examples:**

        Input: "I received an email from customer ABNMST saying that his payment is due next week. Can you predict his payment date and draft a reminder email?"
        Available Tools: predict_payment_timing, draft_email, summarize_text, multiply_numbers, divide_numbers, calculate_sum, convert_to_bullets
        Output: ["Predict payment date for customer ABNMST", "Draft a reminder email for customer ABNMST"]

        Input: "My computer is running very slow and I keep getting a blue screen error. Can you help?"
        Available Tools: troubleshoot_computer, run_diagnostics, restart_system
        Output: ["Troubleshoot slow computer performance", "Troubleshoot blue screen error"]

        Input: "Summarize document.txt and translate it to Spanish"
        Available Tools: summarize_text, translate_text, extract_text
        Output: ["Summarize document.txt", "Translate the summary to Spanish"]

        **Output (JSON array of strings):**
    """

        arguments_str = self.llm_call(prompt, llm_type)

        if not arguments_str:
            logging.error("LLM returned an empty response.")
            return []

        
        try:
            inferred_goals = json.loads(arguments_str)
            if not isinstance(inferred_goals, list):
                logging.error("LLM did not return a list of goals in JSON format.")
                return []
            
            # Additional check to ensure each item in the list is a string
            for goal in inferred_goals:
                if not isinstance(goal, str):
                    logging.error("LLM returned a non-string goal.")
                    return []
                
            return inferred_goals

        except json.JSONDecodeError as e:
            logging.error(f"Error parsing inferred goals as JSON: {e}")
            logging.error(f"Problematic JSON string: {arguments_str}")
            return []

    async def _generate_graph_from_goal(self, goal: str, tools: Dict, llm_type: str, dynamic_example_selector, instruction_retriever, selected_tools: List[Dict]=None,selected_agents: List[str]=None, retry_count: int = 0, max_retries: int = 3,feedback_prompt_json_error: str = None,feedback_from_AI_verifier: str = None,previous_graph: Dict = None,previous_context: str = None) -> Dict:
        """
        Uses the LLM to convert the user's goal into a graph representation (adjacency list).
        Now generates simpler node names and stores entity information as attributes.
        """
        logging.info(f"Generating plan for goal: {goal} with llm_type: {llm_type}")
        client, model_name = self.get_client_and_model(llm_type)

        print("selected_agents......", selected_agents)
        # ---  Use selected_tools descriptions if available ---
        if selected_tools:
            tool_descriptions = "\n".join([f"{tool['name']}: {tool['description']}" for tool in selected_tools])
            print(" tool_desc due to select--->", tool_descriptions)
        else:
            tool_descriptions = _get_tool_descriptions(tools,selected_agents=selected_agents)
        # --- End of modification ---

        logging.info(f"selected tools based on Agent {tool_descriptions}")

        relevant_instructions = instruction_retriever.get_relevant_instructions(goal, num_instructions=2, similarity_threshold=0.7, api_key=self.gemini_api_key)
        print (" Relevant Instrucitons--->", relevant_instructions)

        # Format instructions for the prompt, handling missing fields
        instructions_str = ""
        if relevant_instructions:
            for instruction in relevant_instructions:
                instructions_str += "**Instruction:**\n"
                if "Input" in instruction:
                    instructions_str += f"  Input: {instruction['Input']}\n"
                if "Goal" in instruction:
                    instructions_str += f"  Goal: {instruction['Goal']}\n"
                if "Tools to be used" in instruction:
                    instructions_str += f"  Tools to be used: {', '.join(instruction['Tools to be used'])}\n"
                if "Sequence of Tools" in instruction:
                    logging.info(f"GENERATE_GRAPH_FROM_GOAL: Sequence of Tools Type: {type(instruction['Sequence of Tools'])}, Value: {instruction['Sequence of Tools']}")
                    instructions_str += f"  Sequence of Tools: {', '.join(instruction['Sequence of Tools'])}\n"
                if "Graph" in instruction:
                    graph_json = instruction["Graph"] # Use the dictionary directly
                    instructions_str += f"  Relevant Graph:\n```json\n{json.dumps(graph_json, indent=2)}\n```\n" 
                    #try: # Handle potential JSON parsing errors for Graph
                        #graph_json = json.loads(instruction["Graph"]) # Parse JSON string to object
                        #instructions_str += f"  Relevant Graph:\n```json\n{json.dumps(graph_json, indent=2)}\n```\n" # Format JSON and add code block
                    #except (json.JSONDecodeError, TypeError) as e:
                        #logging.warning(f"Warning: Could not parse or format 'Graph' from instructions: {e}")
                        #instructions_str += f"  Relevant Graph: {', '.join(instruction['Graph'])}\n"

                if "User feedback for the plan" in instruction: # New: User feedback
                    instructions_str += f"  User feedback for the plan: {instruction['User feedback for the plan']}\n" # New: User feedback
                if "Revised Tools" in instruction: # New: Revised Tools
                    instructions_str += f"  Revised Tools: {instruction['Revised Tools']}\n" # New: Revised Tools
                if "Revised Sequence of Tools" in instruction: # New: Revised Sequence
                    instructions_str += f"  Revised Sequence of Tools: {instruction['Revised Sequence of Tools']}\n" # New: Revised Sequence
                    
                if instruction.get("Additional Information"):
                    instructions_str += f"  Additional Information: {instruction['Additional Information']}\n"
                instructions_str += "---\n"
        print (" Relevant Instrucitons II--->", relevant_instructions)

        # --- Dynamic Example Insertion ---
        # (Assuming you still want to use dynamic example selection)
        """examples_str = dynamic_example_selector.insert_examples_into_prompt(
            prompt="",
            goal=goal,
            example_type="_generate_graph_from_goal",
            num_good_examples=2,
            num_bad_examples=0,
            similarity_threshold=0.2,
            api_key=self.gemini_api_key # Pass api_key here
            )"""

        # ---  Include the selected tools in the prompt ---
        selected_tools_str = ""
        if selected_tools:
            selected_tools_str = "\n".join([f"- {tool['name']}: {tool['description']}" for tool in selected_tools])

        prompt = f"""
        You are an expert AI planner that **ONLY outputs plans in a specific JSON format**. Your task is to convert user goals into a directed graph representing the sequence of actions and their dependencies.

        **IMPORTANT: You MUST represent the graph as a JSON ADJACENCY LIST.**

        **DO NOT use "nodes" and "edges" arrays.**

        **DO NOT include any descriptive text or explanations outside the JSON.**

        **Reasoning Process (Think Step-by-Step):**
        1. **Identify the Initial Action:** What is the first tool or action that needs to be performed based on the goal?
        2. **Determine Necessary Inputs:** What are the inputs required for the initial action? Are these available directly from the goal or do they need to be computed by a previous step?
        3. **Identify Subsequent Actions:** Based on the output of the initial action and the overall goal, what are the next logical tools or actions required?
        4. **Establish Dependencies:**  For each subsequent action, determine which preceding action(s) provide the necessary inputs.
        5. **Map Inputs and Outputs:** Ensure the output of one action is correctly used as the input for the next dependent action. Use the `{{node_name}}` placeholder for this.
        6. **Consider All Available Tools:**  Evaluate which of the provided tools are relevant and necessary to achieve the goal.
        7. **Handle Edge Cases and Constraints:**  Are there any specific instructions or constraints that need to be considered during the planning process?
        8. **Structure the Graph:**  Organize the actions and dependencies into the specified JSON format, ensuring the "start" and "end" nodes are correctly placed.


        **Output Format: JSON Adjacency List -  CRITICAL NODE STRUCTURES**

        The graph MUST be represented using a JSON dictionary where:

        1.  **Keys are NODE NAMES (strings).**  Node names should be simple identifiers (e.g., "start", "tool_node_1", "tool_node_2", "end").

        2.  **Values for each key are NODE DEFINITIONS, and they MUST adhere to the following STRUCTURES:**

            *   **"start" NODE - STRUCTURE MUST BE:  `"start": ["node_name_1", "node_name_2", ...]`** 
                The value for the "start" node **MUST BE a JSON ARRAY of strings**, listing the names of the starting nodes of the plan. It is a **JSON LIST/ARRAY**, not a single string.
                Example: `"start": ["identify_relevant_pkl_files_1", "summarize_text_1"]` OR if only one starting node: `"start": ["multiply_numbers_1"]`

            *   **TOOL NODES (representing tools to be executed) - STRUCTURE MUST BE: (JSON Dictionary)**
                ```json
                "tool_node_1": {{
                    "tool": "tool_name",  // **REQUIRED**: String, name of the tool to execute
                    "input": {{...}},     // **REQUIRED**: JSON dictionary, input arguments for the tool
                    "next": [...]       // **REQUIRED for tool nodes**: JSON array of strings (node names), indicating next nodes to execute after this tool
                }}
                ```
                Example:
                ```json
                "tool_node_1": {{
                    "tool": "summarize_text",
                    "input": {{"text": "document_content"}},
                    "next": ["tool_node_2"]
                }}
                ```

            *   **"end" NODE - STRUCTURE MUST BE: `"end": {{}}` (Empty JSON Dictionary)**
                The value for the "end" node **MUST be an EMPTY JSON DICTIONARY** -  `{{}}`. It is a **JSON OBJECT/DICTIONARY**, not a string.
                Example: `"end": {{}}`


        **Example of Correct JSON Adjacency List Graph Format (Correct "start" and "end" Node Structures):**

        ```json
        {{
            "start": ["identify_relevant_pkl_files_1"], 
            "identify_relevant_pkl_files_1": {{
                "tool": "identify_relevant_pkl_files",
                "input" : {{"goal": "find contracts for elandia"}},
                "next": ["answer_questions_1"]
            }},
            "answer_questions_1": {{
                "tool": "answer_questions",
                "input" : {{"file_names": "{{identify_relevant_pkl_files_1}}", "goal": "what is the contract value?"}},
                "next": ["end"]
            }},
            "end": {{}}  // Correct "end" node structure - empty dictionary
        }}
        ```

        **IMPORTANT - Correct "input" Structure for Tool Nodes:**

        For every TOOL NODE, the `"input"` key MUST have a **JSON DICTIONARY** as its value. 

        **INCORRECT "input" (String - DO NOT DO THIS):**
        ```json
        "tool_node_x": {{
            "tool": "some_tool",
            "input": "This is WRONG! Input must be a dictionary, not a string!", 
            "next": [...] 
        }}
        ```

        **CORRECT "input" (JSON Dictionary - DO THIS):**
        ```json
        "tool_node_x": {{
            "tool": "some_tool",
            "input": {{ 
                "argument_name_1": "argument_value_1",
                "argument_name_2": 123 
            }},
            "next": [...]
        }}
        ```

        **Key Instructions - PLEASE FOLLOW THESE EXACTLY:**

        *   **Output MUST be ONLY VALID JSON.**
        *   **Use the ADJACENCY LIST format as described above.**
        *   **Include "start" and "end" nodes.**
        *   **For each tool node, include "tool", "input", and "next" keys.**
        *   **Represent dependencies using node names and the `next` arrays.**
        *   **Do NOT include "nodes" and "edges" arrays.**
        *   **Do NOT include any descriptions or text outside the JSON structure.**

        **Available Tools and their Descriptions:**
        {tool_descriptions}

        **Recommended list of Tools and their Order for this Goal as confirmed by User. If this is available, please use this as the list, dont change the order or tools:**
        {selected_tools_str}

        **Important Instructions:**
        - Use the provided tool descriptions to understand the purpose and usage of each tool.
        - If there are recommended tools, use that list preferably.
        - Generate simple node names that reflect the tool being used (e.g., `multiply_numbers`, `divide_numbers`, `draft_email`).
        - Do not embed entity information directly in the node names.
        - Store entity information (e.g., customer names, amounts, recipients) as attributes or properties of the nodes.
        - Ensure that the graph accurately reflects the dependencies between tools based on the goal and the tool descriptions.
        - The graph should always start with a "start" node and end with an "end" node.
        - **Crucially, when a tool's input depends on the output of a previous tool, use the format `{{node_name}}` within the `input` dictionary to represent that dependency. For example, if `draft_email` needs the output of `predict_payment_timing_1`, the input for `draft_email` might look like: `{{'recipient': '{{get_account_emails_1}}', 'content': 'Predicted day: {{predict_payment_timing_1}}'}}`.**
        - **Your output must be a valid JSON string in the specified format and nothing else. Do not include any introductory or explanatory text.**
        **The value of the "start" key MUST ALWAYS be a JSON array (a list), even if there is only one starting node.  For example:**

        **CORRECT (One Starting Node):**
        "start": ["node_1"]

        **INCORRECT (One Starting Node - DO NOT DO THIS):**
        "start": "node_1"  // WRONG - Must be a list!
        "start": {{"tool": "start", "next": "node_1"}} // WRONG - Must be a list of node names!

        **CRITICAL:** EVERY tool node MUST have a `"tool"` key, and its value MUST be the name of an available tool.  Do NOT omit the `"tool"` key. Do NOT set the `"tool"` key to null or an empty string.

          **Example (CORRECT):**

        ```json
        {{
            "start": ["identify_files_1"],
            "identify_files_1": {{
                "tool": "identify_relevant_pkl_files",  <-- CORRECT
                "input": {{
                    "goal": "Find all documents about project X"
                }},
                "next": ["summarize_1"]
            }},
            "summarize_1": {{
                "tool": "summarize_text",  <-- CORRECT
                "input": {{
                    "text": "{{identify_files_1}}"
                }},
                "next": ["end"]
            }},
            "end": {{}}
        }}
        ```
        **Example (INCORRECT - DO NOT DO THIS):**

        ```json
        {{
            "start": ["identify_files_1"],
            "identify_files_1": {{
                "input": {{  # <--- MISSING "tool" key!
                    "goal": "Find all documents about project X"
                }},
                "next": ["summarize_1"]
            }},
            "summarize_1": {{
                "tool": null,   <--  "tool" key present, but value is null/None - INCORRECT
                "input": {{
                    "text": "{{identify_files_1}}"
                }},
                "next": ["end"]
            }},
            "end": {{}}
        }}
        ```
        **Relevant Instructions for the goal:**
        {instructions_str}

        Goal: "{goal}"

        **Previous Conversation Context (if any):**
        {previous_context if previous_context else "No prior conversation."}

        Graph:
        """

        # Add feedback to the prompt IF it's provided.
        if feedback_from_AI_verifier:
            prompt += f"""
        Previous attempt to generate a plan for this goal resulted in the following feedback from an AI verifier:

        **AI Verifier Feedback:**
        {feedback_from_AI_verifier}

        **Previous Graph for which the feedback was given by AI verifier:**
        ```json
        {json.dumps(previous_graph, indent=4)}
        ```

        Please use this feedback to correct the plan. Pay close attention to the specific issues
        identified by the verifier.
        """

        prompt += """
        Graph:
        """


        if feedback_prompt_json_error:
            prompt += f"""

        **Previous response was NOT valid JSON and caused a parsing error. Please generate a VALID JSON graph response, taking into account the following feedback:**
        **Feedback on JSON Error:**
        {feedback_prompt_json_error}

        **Ensure the Revised Graph is Valid JSON:**
        """

        logging.info(f"First prompt: {prompt}")
        graph_str = self.llm_call(prompt, llm_type)
        try:
            print("graph before json load--->", graph_str)
            graph = json.loads(graph_str)
            print("graph before build--->", graph)
            #graph = build_graph(graph)
            print("graph after build--->", graph)
            return graph

        except json.JSONDecodeError as e:
            invalid_json_string = graph_str # Capture invalid JSON string
            logging.error(f"Invalid JSON returned from LLM (Attempt {retry_count + 1}): {e}")
            logging.error(f"Problematic JSON string: {invalid_json_string}")

            if retry_count < max_retries:
                feedback_prompt_json_error = f"Your response was expected to be in valid JSON format for a graph, but it failed to parse with error: '{e}'. The invalid JSON string was: '{invalid_json_string}'. Please return ONLY valid JSON graph in your next response."
                logging.warning(f"Regenerating graph due to JSON parsing error. (Retry {retry_count + 1}/{max_retries})")
                return await self._generate_graph_from_goal( # Recursive retry for graph generation
                    goal, tools, llm_type, dynamic_example_selector, instruction_retriever, selected_tools, selected_agents, retry_count + 1, max_retries, feedback_prompt_json_error # Increment retry count, pass feedback
                )
            else:
                logging.error(f"Maximum retries exceeded for graph regeneration after JSON parsing errors.")
                logging.error(f"Original JSON parsing error: {e}")
                logging.error(f"Last invalid JSON String: {invalid_json_string}")
                return {} # Return empty graph or handle error as needed    
     
    
    async def _generate_graph_from_goal_old(self, goal: str, tools: Dict, llm_type: str, dynamic_example_selector, instruction_retriever, selected_tools: List[Dict]=None,selected_agents: List[str]=None, retry_count: int = 0, max_retries: int = 3,feedback_prompt_json_error: str = None) -> Dict:
        """
        Uses the LLM to convert the user's goal into a graph representation (adjacency list).
        Now generates simpler node names and stores entity information as attributes.
        """
        logging.info(f"Generating plan for goal: {goal} with llm_type: {llm_type}")
        client, model_name = self.get_client_and_model(llm_type)

        print("selected_agents......", selected_agents)
        # ---  Use selected_tools descriptions if available ---
        if selected_tools:
            tool_descriptions = "\n".join([f"{tool['name']}: {tool['description']}" for tool in selected_tools])
            print(" tool_desc due to select--->", tool_descriptions)
        else:
            tool_descriptions = _get_tool_descriptions(tools,selected_agents=selected_agents)
        # --- End of modification ---

        logging.info(f"selected tools based on Agent {tool_descriptions}")

        relevant_instructions = instruction_retriever.get_relevant_instructions(goal, num_instructions=2, similarity_threshold=0.7, api_key=self.gemini_api_key)
        print (" Relevant Instrucitons--->", relevant_instructions)

        # Format instructions for the prompt, handling missing fields
        instructions_str = ""
        if relevant_instructions:
            for instruction in relevant_instructions:
                instructions_str += "**Instruction:**\n"
                if "Input" in instruction:
                    instructions_str += f"  Input: {instruction['Input']}\n"
                if "Goal" in instruction:
                    instructions_str += f"  Goal: {instruction['Goal']}\n"
                if "Tools" in instruction:
                    instructions_str += f"  Tools to be used: {', '.join(instruction['Tools'])}\n"
                if "Sequence" in instruction:
                    instructions_str += f"  Sequence of Tools: {', '.join(instruction['Sequence'])}\n"
                if "Graph" in instruction:
                    instructions_str += f"  Relevant Graph: {', '.join(instruction['Graph'])}\n"
                if instruction.get("Additional Information"):
                    instructions_str += f"  Additional Information: {instruction['Additional Information']}\n"
                instructions_str += "---\n"
        print (" Relevant Instrucitons II--->", relevant_instructions)

        # --- Dynamic Example Insertion ---
        # (Assuming you still want to use dynamic example selection)
        """examples_str = dynamic_example_selector.insert_examples_into_prompt(
            prompt="",
            goal=goal,
            example_type="_generate_graph_from_goal",
            num_good_examples=2,
            num_bad_examples=0,
            similarity_threshold=0.2,
            api_key=self.gemini_api_key # Pass api_key here
            )"""

        # ---  Include the selected tools in the prompt ---
        selected_tools_str = ""
        if selected_tools:
            selected_tools_str = "\n".join([f"- {tool['name']}: {tool['description']}" for tool in selected_tools])

        prompt_template = f"""
        You are an AI planner that converts user goals into a directed graph representing the sequence of actions and their dependencies.

        Nodes in the graph represent tools or actions to be executed.
        Edges represent dependencies between nodes, indicating the order of execution.

        **Reasoning Process (Think Step-by-Step):**

        1. **Identify the Initial Action:** What is the first tool or action that needs to be performed based on the goal?
        2. **Determine Necessary Inputs:** What are the inputs required for the initial action? Are these available directly from the goal or do they need to be computed by a previous step?
        3. **Identify Subsequent Actions:** Based on the output of the initial action and the overall goal, what are the next logical tools or actions required?
        4. **Establish Dependencies:**  For each subsequent action, determine which preceding action(s) provide the necessary inputs.
        5. **Map Inputs and Outputs:** Ensure the output of one action is correctly used as the input for the next dependent action. Use the `{{node_name}}` placeholder for this.
        6. **Consider All Available Tools:**  Evaluate which of the provided tools are relevant and necessary to achieve the goal.
        7. **Handle Edge Cases and Constraints:**  Are there any specific instructions or constraints that need to be considered during the planning process?
        8. **Structure the Graph:**  Organize the actions and dependencies into the specified JSON format, ensuring the "start" and "end" nodes are correctly placed.

        **Available Tools and their Descriptions:**
        {tool_descriptions}

        **Recommended list of Tools and their Order for this Goal as confirmed by User. If this is available, please use this as the list, dont change the order or tools:**
        {selected_tools_str}

        **Important Instructions:**
        - Use the provided tool descriptions to understand the purpose and usage of each tool.
        - If there are recommended tools, use that list preferably. 
        - Generate simple node names that reflect the tool being used (e.g., `multiply_numbers`, `divide_numbers`, `draft_email`).
        - Do not embed entity information directly in the node names.
        - Store entity information (e.g., customer names, amounts, recipients) as attributes or properties of the nodes.
        - Ensure that the graph accurately reflects the dependencies between tools based on the goal and the tool descriptions.
        - The graph should always start with a "start" node and end with an "end" node.
        - If the goal involves performing calculations and using the results in subsequent steps (like summarization or drafting emails), make sure the calculation nodes are executed *before* the nodes that depend on them.
        - **Crucially, when a tool's input depends on the output of a previous tool, use the format `{{node_name}}` within the `input` dictionary to represent that dependency. For example, if `draft_email` needs the output of `predict_payment_timing_1`, the input for `draft_email` might look like: `{{'recipient': '{{get_account_emails_1}}', 'content': 'Predicted day: {{predict_payment_timing_1}}'}}`.**
        - **Your output must be a valid JSON string in the specified format and nothing else. Do not include any introductory or explanatory text.**

            
        Represent the graph using a JSON format where each node can have attributes:
        {{
            "start": ["multiply_numbers_1"],
            "multiply_numbers_1": {{
                "tool": "multiply_numbers",
                "input" : {{"num1": 3, "num2": 3}},
                "next": ["divide_numbers_1"]
            }},
            "divide_numbers_1": {{
                "tool": "divide_numbers",
                "input" : {{"num1": "{{multiply_numbers_1}}", "num2": 4}},
                "next": ["draft_email_1"]
            }},
            "draft_email_1": {{
                "tool": "draft_email",
                "input": {{"recipient": "ddd@dd.com", "subject": "Calculation Result", "content": "The result of the calculation is {{divide_numbers_1}}"}},
                "next": ["end"]
            }},
            "end": {{}}
        }}
            """

        prompt = f"""
        {prompt_template}

        **Relevant Instructions for the goal:**
        {instructions_str}
            

        Goal: "{goal}"

        Graph:
        """

        if feedback_prompt_json_error:
            prompt += f"""

        **Previous response was NOT valid JSON and caused a parsing error. Please generate a VALID JSON graph response, taking into account the following feedback:**
        **Feedback on JSON Error:**
        {feedback_prompt_json_error}

        **Ensure the Revised Graph is Valid JSON:**
        """

        print(f"llm type: {llm_type}")
        graph_str = self.llm_call(prompt, llm_type)
        try:
            graph = json.loads(graph_str)
            graph = build_graph(graph)
            return graph

        except json.JSONDecodeError as e:
            invalid_json_string = graph_str # Capture invalid JSON string
            logging.error(f"Invalid JSON returned from LLM (Attempt {retry_count + 1}): {e}")
            logging.error(f"Problematic JSON string: {invalid_json_string}")

            if retry_count < max_retries:
                feedback_prompt_json_error = f"Your response was expected to be in valid JSON format for a graph, but it failed to parse with error: '{e}'. The invalid JSON string was: '{invalid_json_string}'. Please return ONLY valid JSON graph in your next response."
                logging.warning(f"Regenerating graph due to JSON parsing error. (Retry {retry_count + 1}/{max_retries})")
                return await self._generate_graph_from_goal( # Recursive retry for graph generation
                    goal, tools, llm_type, dynamic_example_selector, instruction_retriever, selected_tools, selected_agents, retry_count + 1, max_retries, feedback_prompt_json_error # Increment retry count, pass feedback
                )
            else:
                logging.error(f"Maximum retries exceeded for graph regeneration after JSON parsing errors.")
                logging.error(f"Original JSON parsing error: {e}")
                logging.error(f"Last invalid JSON String: {invalid_json_string}")
                return {} # Return empty graph or handle error as needed  
    
    def _error_to_feedback_prompt(self, error: Exception, tool_name: str, arguments: Dict, goal: str, memory_manager:MemoryManager, tool_descriptions: str) -> str:
        """
        Converts an error message into a feedback prompt for argument regeneration using an LLM.

        Args:
            error: The exception object.
            tool_name: The name of the tool that caused the error.
            arguments: The arguments that were passed to the tool.
            llm_client: The LLM client (AzureOpenAI or genai.GenerativeModel).
            model_name: The name of the LLM model to use.
            goal: The original goal of the plan.
            memory: The current memory of the agent.
            tool_descriptions: A string containing the descriptions of all available tools.

        Returns:
            A feedback prompt string.
        """
        error_message = str(error)

        # Log the error with relevant information for debugging
        logging.error(f"Error during execution of tool '{tool_name}'")
        logging.error(f"Error message: {error_message}")
        logging.error(f"Arguments: {arguments}")
        logging.error(f"Goal: {goal}")

        prompt = f"""
        You are an AI assistant that helps in debugging and providing feedback for an agent framework.

        An error occurred during the execution of the tool '{tool_name}'.

        Original Goal: {goal}

        Current Memory: {memory_manager.results}

        Available Tools: 
        {tool_descriptions}

        Error message: {error_message}

        The arguments provided to the tool were: {arguments}

        Please provide a concise and informative feedback prompt that can be used to regenerate the arguments for this tool. The feedback should focus on what went wrong and what should be corrected, taking into account the original goal, current memory and available tools.

        Feedback:
        """
        
        feedback_prompt = self.llm_call(prompt, "gemini")
        
    
        # Log the generated feedback prompt
        logging.info(f"Generated feedback prompt: {feedback_prompt}")

        return feedback_prompt
    
    
    def _verify_graph_with_llm(self, goal: str, graph: Dict, tools: Dict, llm_type: str = None, selected_agents: List[str] = None) -> Union[bool, str]:
        """
        Verifies if the generated graph matches the goal using an LLM.

        Args:
            goal: The original user goal.
            graph: The generated graph as a dictionary.
            tools: A dictionary of available tools.
            llm_type: The type of LLM being used.
            selected_agents: List of agents.

        Returns:
            True if the graph is valid and matches the goal, otherwise a string with feedback.
        """
        client, model_name = self.get_client_and_model(llm_type)
        tool_descriptions = _get_tool_descriptions(tools, selected_agents=selected_agents)

        prompt = f"""
        You are an AI graph verifier. Your task is to analyze a proposed plan (represented as a JSON graph)
        and determine if it is valid and likely to achieve the user's goal.

        **User Goal:** {goal}

        **Proposed Graph (JSON):**
        ```json
        {json.dumps(graph, indent=4)}
        
        Available Tools:
        {tool_descriptions}

        Graph Format (Adjacency List):

        The graph MUST be a JSON object where:

        Keys: Node names (strings). These should be unique and descriptive (e.g., "identify_files_1", "summarize_text_1").
        Values: Node definitions (JSON objects).
        Node Definition Format:

        Each node (except "start" and "end") MUST have the following structure:

        {{
            "tool": "tool_name",  // REQUIRED.  The name of the tool to use.
            "input": {{...}},    // REQUIRED.  A JSON object containing the input arguments for the tool.
            "next": [...]        // REQUIRED.  A JSON *array* of node names to execute *after* this node.
        }}
        
        Special Nodes:

        "start": The value of the "start" node MUST be a JSON array of strings, listing the starting node(s). Example: "start": ["node_1", "node_2"]
        "end": The value of the "end" node MUST be an empty JSON object: "end": {{}}
        
        CRITICAL RULES (MUST FOLLOW):

        Valid JSON: The entire graph MUST be valid JSON.
        Adjacency List: The graph MUST use the adjacency list format described above. Do NOT use "nodes" and "edges" arrays.
        "next" Key: Dependencies between nodes MUST be represented using the "next" key. The value of "next" MUST be a JSON array (list) of node names. Do NOT use any other key (like "output") to represent dependencies.
        "start" Node: The value of "start" MUST be a JSON array.
        "end" Node: The value of "end" MUST be an empty JSON object {{}}.
        Tool Names: Tool names in the "tool" field MUST be one of the tools listed in "Available Tools".
        Input: The input should strictly follow the arguments mentioned in the tools description.
        All Nodes Reachable: Check if there are any orphaned or isolated nodes/subgraphs that aren't connected to the rest of the graph.
        Verification Tasks:

        Format Check: Check that the graph strictly adheres to the JSON format described above, paying close attention to "start", "end", "next", and node structure. Report specific format violations. For example:
        "The 'start' node's value is not a list."
        "Node 'xyz' is missing the 'next' key."
        "Node 'abc' has an invalid 'next' value (not a list)."
        "The graph is not valid JSON."
        "Node 'xyz' uses an unknown tool: 'invalid_tool'."
        "Node 'abc' is missing the required 'input' key."
        "Node 'def' has a 'next' value that is not a list: 'some_string'."
        Dependency Check: Verify that all nodes listed in a "next" array actually exist in the graph.
        Tool Usage Check: Verify that each node uses a valid tool name from the "Available Tools" list, and that the input arguments are appropriate for that tool (based on the descriptions).
        Logical Flow (High Level): Does the overall sequence of tools seem reasonable for achieving the goal?
        Output:

        Return a single JSON object with the following format:

        {{
            "score": 0.9,  // A floating-point number between 0.0 and 1.0 (inclusive).
            "feedback": "The plan is mostly correct, but..." // A string describing any issues.
        }}
        
        score: A floating-point number between 0.0 and 1.0 (inclusive), representing the overall quality of the graph. 1.0 means perfect; 0.0 means completely incorrect.
        feedback: A string providing specific, concise feedback about any problems found. If the score is 1.0, the feedback can be "The graph is valid and well-formed." or similar.
       
         Scoring Guidelines:

        1.0: Perfect. The graph is valid JSON, follows all format rules, uses tools correctly, and represents a logically sound plan to achieve the goal.
        0.8 - 0.99: Minor issues. The graph is mostly correct, but has some minor formatting errors, slightly suboptimal tool choices, or minor logical inconsistencies.
        0.5 - 0.79: Moderate issues. The graph has significant format errors, tool misuse, or logical flaws, but still attempts to address the goal.
        0.2 - 0.49: Major issues. The graph has severe format errors, uses tools inappropriately, or has a fundamentally flawed logical structure.
        0.0 - 0.19: Completely incorrect. The graph is invalid JSON, bears no resemblance to the required format, or is completely unrelated to the goal.
        
        Be strict about format. Output ONLY the JSON, and nothing else. Do NOT include any introductory or explanatory text. Do not include markdown formatting like ```json.

        """

        response_str = self.llm_call(prompt, llm_type)

        try:
            response_json = json.loads(response_str)
            score = float(response_json["score"])  # Ensure it's a float
            feedback = response_json["feedback"]

            if 0.0 <= score <= 1.0:
                return score, feedback # Return both score and feedback
            else:
                logging.error(f"LLM returned an invalid score: {score}")
            return 0.0, f"Invalid score from LLM: {score}. Response: {response_str}" # return 0 and feedback

        except (json.JSONDecodeError, KeyError, ValueError) as e:
            logging.error(f"Error parsing LLM response for graph verification: {e} Response: {response_str}")
            return 0.0, f"Invalid response from LLM: {response_str}"  # Return 0 and the raw response

    def _verify_tool_result_with_llm(self, tool_name: str, tool_description: str, arguments: Dict, result: Any, goal: str, memory_manager: MemoryManager, llm_type: str = None) -> Union[bool, str]:
        """
        Verifies if the result of a tool execution is valid and makes sense in the context of the goal using an LLM.

        Args:
            tool_name: The name of the tool.
            tool_description: Description of the tool.
            arguments: Arguments passed to the tool.
            result: The output of the tool execution.
            goal: The user's goal.
            memory: Current memory state.
            llm_type: LLM type.

        Returns:
            True if the result is valid, or feedback string if invalid.
        """
        client, model_name = self.get_client_and_model(llm_type)
        
        prompt = f"""
        You are an AI tool result verifier. Your task is to determine if the result of executing the tool '{tool_name}' is valid and reasonable in the context of the user's goal and tool description.

        **Tool Name:** {tool_name}
        **Tool Description:** {tool_description}
        **Arguments Used:** {json.dumps(arguments)}
        **Tool Result:** {result}
        **User Goal:** {goal}
        **Current Memory:** {json.dumps(memory_manager.results)}

        **Verification Tasks:**

        1. **Plausibility Check:** Does the result seem plausible given the tool description and arguments?
        2. **Goal Alignment:** Does the result contribute towards achieving the user's goal?
        3. **Type and Format Validation:** Is the result in the expected format and data type for this tool? (e.g., is it a number when a number is expected, is it text if text is expected?)
        4. **Error Detection:** Does the result indicate an error or failure in the tool execution (even if no exception was raised)?

        **Instructions:**

        - If the result is valid and reasonable, return "True".
        - If the result is invalid, unreasonable, or indicates an error, return a concise feedback string explaining the issue.  Be specific about what makes the result invalid.
        - Focus on actionable feedback that could help regenerate arguments or revise the plan in subsequent steps.
        - Do not provide any introductory or explanatory text, just "True" or the feedback string.

        **Example:**

        Tool Name: predict_payment_timing
        Tool Description: Predicts the payment timing for a customer. Returns a date string.
        Arguments Used: {{"customer_name": "Acme Corp", "amount": 10000}}
        Tool Result: "Invalid Customer ID"
        User Goal: Predict payment date for Acme Corp.

        Verification Result:
        "The tool returned an error message 'Invalid Customer ID', indicating a failure. The result is not valid."

        **Return your verification result:**
        """

        verification_result = self.llm_call(prompt, "gemini_reasoning")

        if verification_result.lower() == "true":
            return True
        else:
            return verification_result
        
    
    def _verify_final_result_with_llm(self, goal: str, results: list, llm_type: str = None) -> Union[float, str]:
        """
        Verifies if the final result of plan execution meets the user's goal,
        returning a score between 0.0 and 1.0, and feedback.
        """
        client, model_name = self.get_client_and_model(llm_type)

        prompt = f"""
            You are an AI result verifier.  Your task is to assess whether the results of a
            plan execution successfully achieve the user's original goal.

            **User Goal:** {goal}

            **Execution Results:**
            ```json
            {json.dumps(results, indent=4)}
            Use code with caution.
            Python
            Instructions:

            Carefully compare the Execution Results with the User Goal.
            Determine if the results, taken as a whole, completely and satisfactorily
            address the goal.
            Return a JSON object with the following format:
            {{
                "score": 0.9,  // A floating-point number between 0.0 and 1.0 (inclusive).
                "feedback": "The plan is mostly correct, but..." // String: specific issues.
            }}
            Use code with caution.
            Json
            score: A floating-point number between 0.0 and 1.0 (inclusive),
            representing the overall quality of how well the results achieve the goal.
            1.0 means perfect; 0.0 means completely failed.
            feedback: A string providing specific, concise feedback about any
            problems found. If the score is 1.0, the feedback can be something like
            "The results completely satisfy the goal." or "The goal was achieved
            successfully." Always include the feedback string, even with a
            score of 1.0.
            Scoring Guidelines:

            1.0: Perfect. The results completely and correctly address the user's goal.
            0.8 - 0.99: Minor issues. The results mostly address the goal, but there are some minor omissions, inaccuracies, or inefficiencies.
            0.5 - 0.79: Moderate issues. The results partially address the goal, but there are significant omissions, inaccuracies, or inefficiencies.
            0.2 - 0.49: Major issues. The results largely fail to address the goal, or contain major errors.
            0.0 - 0.19: Completely incorrect. The results are unrelated to the goal or completely wrong.
            Important: Return ONLY the JSON, and nothing else and DO NOT include any markdown, like ```.
            Output:
            """
        response_str = self.llm_call(prompt, llm_type)

        try:
            response_json = json.loads(response_str)
            score = float(response_json["score"])
            feedback = response_json["feedback"]

            if 0.0 <= score <= 1.0:
                return score, feedback  # Return both score and feedback
            else:
                logging.error(f"LLM returned an invalid score: {score}")
                return 0.0, f"Invalid score from LLM: {score}. Response: {response_str}"

        except (json.JSONDecodeError, KeyError, ValueError) as e:
            logging.error(f"Error parsing LLM response for final result verification: {e}. Response: {response_str}")
            return 0.0, f"Invalid response from LLM: {response_str}"
            
    async def _regenerate_graph_with_feedback(self, goal: str, tools: Dict, feedback_plan: str, previous_graph: Dict, llm_type: str = None, selected_agents: List[str] = None) -> Dict:
        """
        Uses the LLM to generate a revised graph based on user feedback.
        """
        client, model_name = self.get_client_and_model(llm_type)

        tool_descriptions = _get_tool_descriptions(tools,selected_agents=selected_agents)

        prompt = f"""
        You are an AI planner that converts user goals into a directed graph representing the sequence of actions and their dependencies.
        You are now tasked with REVISING a previously generated graph based on user feedback.

        Nodes in the graph represent tools or actions to be executed.
        Edges represent dependencies between nodes, indicating the order of execution.

        **Available Tools and their Descriptions:**
        {tool_descriptions}

        **Original Goal:** {goal}

        **Previous Graph:**
        {json.dumps(previous_graph)}

        **User Feedback:** {feedback_plan}

        **Important Instructions:**
        - Use the provided tool descriptions to understand the purpose and usage of each tool.
        - Generate node names that are informative and reflect the tool being used and any relevant entities or parameters. You can use the format `toolname_entity1_entity2` (e.g., `predict_payment_timing_BOFAMAST`, `draft_email_ddd@www.com`).
        - When a tool needs to be applied to multiple entities, create separate nodes for each entity.
        - Ensure that the graph accurately reflects the dependencies between tools based on the goal and the tool descriptions.
        - Do not hardcode any specific tool names or arguments in your reasoning. Rely solely on the provided tool descriptions and the user's feedback.
        - The graph should always start with a "start" node and end with an "end" node.
        - **Your output must be a valid JSON string in the specified adjacency list format and nothing else. Do not include any introductory or explanatory text. Do not include any labels or keys other than the node names.**

        Represent the graph using an adjacency list format:
        {{
            "start": ["node_name_1"],
            "node_name_1": ["node_name_2", "node_name_3"],
            "node_name_2": ["end"],
            "node_name_3": ["end"]
        }}

        **Revised Graph (based on feedback):**
        """
        graph_str = self.llm_call(prompt, llm_type)

        try:
            revised_graph = json.loads(graph_str)
            return revised_graph
        except json.JSONDecodeError as e:
           logging.error(f"Failed to parse JSON response from LLM while generating revised graph due to error {e}.")
           logging.info(f"The invalid JSON string was: {graph_str}.")
           return {}
        
    
    async def _regenerate_graph_with_AI_feedback(self, goal: str, tools: Dict, feedback_plan: str, previous_graph: Dict,  memory_manager: MemoryManager, llm_type: str = None, selected_agents: List[str] = None) -> Dict:
        """
        Uses the LLM to generate a revised graph based on AI feedback (persistent tool failure).
        """

        client, model_name = self.get_client_and_model(llm_type)

        tool_descriptions = _get_tool_descriptions(tools,selected_agents=selected_agents)

        previous_results_str = json.dumps(memory_manager.results, indent=4)

        prompt = f"""
        You are an AI planner that converts user goals into a directed graph representing the sequence of actions and their dependencies.
        You are now tasked with REVISING a previously generated graph based on **AI feedback**, specifically because a tool has been persistently failing **during tool result verification**.

        Nodes in the graph represent tools or actions to be executed.
        Edges represent dependencies between nodes, indicating the order of execution.

        **Available Tools and their Descriptions:**
        {tool_descriptions}

        **Original Goal:** {goal}

        **Previous Graph:**
        {json.dumps(previous_graph)}

         **Previous Execution Results:**
        ```json
        {previous_results_str}
        ```

        **AI Feedback (Persistent Tool Failure during Result Verification):** {feedback_plan}  <--- Clarified feedback type

        **Important Instructions for REVISING the GRAPH (Crucially Different due to Persistent Failure and Result Verification):**

        - **Significantly Re-evaluate the Plan:** The previous plan has proven to be problematic based on *tool result verification*. You need to **think more broadly and creatively** about alternative approaches to achieve the user's goal.
        - **Consider Alternative Tool Sequences:** Explore different sequences of tools that could achieve the same goal, potentially avoiding the tool that is persistently failing ('{{failed_tool_name}}') or using it in a different way.
        - **Analyze Tool Verification Feedback:**  Pay close attention to the **detailed feedback** provided in "**AI Feedback (Persistent Tool Failure during Result Verification):**". This feedback explains *why* the tool's result was deemed invalid. Use this information to guide your plan revision. For example, if the feedback indicates a "File not found" error, consider tools that can locate or create the necessary file, or revise the plan to use a different file source. If the feedback is about incorrect data type, rethink the input data flow to the failing tool.
        - **If Avoiding the Tool is Not Possible:** If the failing tool ('{{failed_tool_name}}') is absolutely essential, then reconsider *when* and *how* it is used in the plan in light of the verification feedback. Maybe it needs different inputs from earlier steps, or needs to be used in a different context, or its arguments need to be prepared in a different way.
        - **Do not just tweak the previous graph.** A minor modification is unlikely to solve a persistent failure issue highlighted by result verification.  **Aim for a substantial revision of the plan structure.**

        **Important Instructions:**
        - Use the provided tool descriptions to understand the purpose and usage of each tool.
        - Generate node names that are informative and reflect the tool being used and any relevant entities or parameters. You can use the format `toolname_entity1_entity2` (e.g., `predict_payment_timing_BOFAMAST`, `draft_email_ddd@www.com`).
        - When a tool needs to be applied to multiple entities, create separate nodes for each entity.
        - Ensure that the graph accurately reflects the dependencies between tools based on the goal and the tool descriptions.
        - Do not hardcode any specific tool names or arguments in your reasoning. Rely solely on the provided tool descriptions and the AI's feedback.
        - The graph should always start with a "start" node and end with an "end" node.
        - **Your output must be a valid JSON string in the specified adjacency list format and nothing else. Do not include any introductory or explanatory text. Do not include any labels or keys other than the node names.**

        Represent the graph using an adjacency list format:
        {{
            "start": ["node_name_1"],
            "node_name_1": ["node_name_2", "node_name_3"],
            "node_name_2": ["end"],
            "node_name_3": ["end"]
        }}

        **Revised Graph (based on AI feedback):**
        """
        graph_str = self.llm_call(prompt, llm_type)

        try:
            revised_graph = json.loads(graph_str)
            return revised_graph
        except json.JSONDecodeError as e:
           logging.error(f"Failed to parse JSON response from LLM while generating revised graph based on AI feedback due to error {e}.")
           logging.info(f"The invalid JSON string was: {graph_str}.")
           return {}
        

    def _verify_arguments_with_llm(self, tool_name: str, tool_description: str, graph: Dict, memory_manager: MemoryManager, arguments: Dict, feedback: str = None, node: str = None, llm_type: str = None) -> Union[bool, str]:
        """
        Verifies the generated arguments using an LLM-based check.

        Args:
            llm_client: The LLM client.
            model_name: The LLM model name.
            tool_name: The name of the tool.
            tool_description: The description of the tool.
            graph: The execution graph.
            memory: The current memory.
            results: The results of previous steps.
            arguments: The generated arguments to verify.
            feedback: Optional user feedback (if available).
            node: The current node in the graph.
            llm_type: The type of LLM.

        Returns:
            True if the arguments are valid, or a string containing feedback if issues are found.
        """
        client, model_name = self.get_client_and_model(llm_type)
        prompt = f"""
        You are an AI verifier that checks the correctness of arguments generated for a tool.

        **Tool Name:** {tool_name}
        **Tool Description:** {tool_description}
        **Graph:** {json.dumps(graph)}
        **Current Memory (Results of Previous Steps):** {json.dumps(memory_manager.results)}
        **Arguments to Verify:** {json.dumps(arguments)}
        **User Feedback (if any):** {feedback if feedback else "None"}

        **Verification Tasks:**

        1. **Check for Completeness:** Ensure all required arguments for the tool (as described in the Tool Description) are present in the `Arguments to Verify`.
        2. **Check for Correct Types:** Verify that the argument values have the correct types (e.g., string, number, etc.) based on the Tool Description.
        3. **Check for Placeholder Substitution:**  If any argument value is a placeholder in the format `{{node_name}}`, verify that it has been correctly replaced with the corresponding result from the `Current Memory`.**Placeholders must be replaced with the actual values, not just mentioned in the text.**
        4. **Check for Consistency with User Feedback:** If user feedback is provided, ensure the arguments are consistent with the feedback.

        **Instructions:**

        - If all arguments are correct, return "True".
        - If any issues are found, return a concise string describing the issues and providing specific suggestions for how to fix them.
        - **Focus on the most critical issues first, especially if placeholder substitutions are incorrect or missing.**
        - **Be prescriptive in your feedback, providing clear instructions on what needs to be changed.**
        - **If a placeholder has not been substituted, explicitly state that it needs to be replaced with the correct value from the Current Memory.**

        **Example:**

        Tool Name: draft_email
        Tool Description: Drafts an email. Arguments: subject (required, string), recipient (required, string), content (required, string).
        Graph: {{"start": ["multiply_numbers_1"], "multiply_numbers_1": ["draft_email_1"], "draft_email_1": {{"tool": "draft_email", "next": ["end"]}}, "end": {{}}}}
        Current Memory: {{"multiply_numbers_1": 6.0}}
        Arguments to Verify: {{"subject": "Calculation Result", "recipient": "test@example.com", "content": "The result is {{multiply_numbers_1}}"}}

        Verification Result:
        "The placeholder `{{multiply_numbers_1}}` in the `content` argument has not been replaced with the actual result. It should be replaced with `6.0` from the Current Memory."

        **Return your verification result:**
        """

        verification_result = self.llm_call(prompt, llm_type)

        if verification_result.lower() == "true":
            return True
        else:
            return verification_result

    def _generate_instruction_content_with_llm(self, request, original_tools: List, original_sequence: List, revised_tools: List = None, revised_sequence: List = None, feedback: str = None, llm_type: str = None) -> str:
        """Generates instruction content using an LLM based on the given parameters."""
        client, model_name = self.get_client_and_model(llm_type)

        prompt = f"""
        You are an AI assistant that generates instruction content for an AI agent based on the following data from a previous execution. The instruction content will be used for future AI tools to learn from, when it encournters with similar User Input or gaol in the future.
        
        **User Input:** {request.unstruct_input if request.unstruct_input else "Not Available"}
        
        **Goal:** {request.goal}
        
        **Original Tools:** {', '.join(original_tools) if original_tools else "Not Available"}
        
        **Original Sequence of Tools:** {', '.join(original_sequence) if original_sequence else "Not Available"}
        
        **User Feedback (if any):** {feedback if feedback else "Not Available"}
        
        **Revised Tools (if any):** {', '.join(revised_tools) if revised_tools else "Not Available"}
        
        **Revised Sequence of Tools (if any):** {', '.join(revised_sequence) if revised_sequence else "Not Available"}
        
        Based on this information, generate a concise and clear instruction content that describes the user input or goal and when you encounter such input or goal articulate the appropriate usage of the tools and the tool sequence. Bacially
        it should articulate the relationship between the User Input or Goal with Tools and Tool Sequence.
        
        If there are any revisions was done based on user feedback, please incorporate the reasoning behind the feedback and the subsequent revisions that were made. The feedback can be correction to rectify a mistake made or it could be 
        a request to add addiitonal tools as a new requirement. so please make sure to understand this difference while creating a generalized knowledge content
        
        Dont include any arguments or specific values in your response. 

        Your output will act as a golden instruction for future AI tools to refer.

        **Output:**
        """
        print('llm call', llm_type)
        instruction_str = self.llm_call(prompt, llm_type)
        return instruction_str
    
    def _select_tools(self, goal: str, tools: Dict, llm_type, selected_agents: List[str]=None) -> List[Dict]:
        """
        Selects the necessary tools for the given goal and returns a list of dictionaries,
        where each dictionary contains the tool name and its description.
        """
        client, model_name = self.get_client_and_model(llm_type)

        # Filter tools based on selected agents
        filtered_tools = {}
        for tool_name, tool_data in tools.items():
            if "agent" in tool_data:
                if "All" in tool_data["agent"]:
                    filtered_tools[tool_name] = tool_data
                elif any(agent in tool_data["agent"] for agent in selected_agents):
                    filtered_tools[tool_name] = tool_data

        prompt = f"""
    You are a helpful AI assistant that selects the best tools to achieve a given goal.

    Goal: {goal}

    Available Tools and Descriptions:
    {_get_tool_descriptions(filtered_tools, selected_agents)}

    Based on the goal, select the necessary tools (without arguments) from the available tools list and return them in the order they should be executed. 
    Only include the tool names, without any arguments or step numbers.

    Return the tools as a list of tool names. For example:

    Goal: draft an email and summarize it
    Tools:
    draft_email
    summarize_text

    Goal: multiply two numbers and divide by 2
    Tools:
    multiply_numbers
    divide_numbers

    Tools:
    """
        selected_tools_str = self.llm_call(prompt, llm_type)

        selected_tools_list = selected_tools_str.split('\n')

        tools_with_descriptions = []
        for tool_name in selected_tools_list:
            tool_name = tool_name.strip()
            if tool_name in filtered_tools:
                tools_with_descriptions.append({
                    "name": tool_name,
                    "description": filtered_tools[tool_name]["description"]
                })

        return tools_with_descriptions
    
    def _revise_tool_selection(self, selected_tools: List[Dict], goal: str, feedback: str, tools: Dict, llm_type: str = None, selected_agents: List[str]=None) -> str:
        """
        Uses the LLM to revise the tool selection based on feedback.
        Returns a string with the revised list of tools.
        """
        client, model_name = self.get_client_and_model(llm_type)
        # 1. Convert selected_tools to a string representation for the prompt:
        selected_tools_str = "\n".join([f"- {tool['name']}: {tool['description']}" for tool in selected_tools])

        # 2. Create the prompt for the LLM:
        prompt = f"""

    AI tool selection assistant. Your goal is to revise a list of initially selected tools based on user feedback.

    Initially Selected Tools:
    {selected_tools_str}

    User feedback:"{feedback}"

    Based on the feedback ,you must do one of the following:
    - **Reordering:** Changing the sequence of tools as per feedback.
    - **Adding Tools:** Including new tools as per feedback.
    - **Removing Tools:** Remove tools as per feedback

    **Output Format:**

    Return the revised list of tools as a newline-separated list of tool names, strictly in the order they should be executed. Only include the tool names. 

    Based on the user feedback, revise the list of selected tools and their order.
    Only include the tool names, without any arguments or step numbers.
    Return the revised list of tools as a list of tool names, in the order they should be executed.

    **Examples:**

    Initially Selected Tools:
    - translate_text: Translates text to a specified language.
    - summarize_text: Summarizes the given text.
    User Feedback: The order of tools is incorrect. The text should be summarized first.
    Revised Tools:
    summarize_text
    translate_text

    Initially Selected Tools:
    - translate_text: Translates text to a specified language.
    - summarize_text: Summarizes the given text.
    User Feedback: Remove translate_text.
    Revised Tools:
    summarize_text

    Initially Selected Tools:
    - summarize_text: Summarizes the given text.
    User Feedback: The tool selection is correct.
    Revised Tools:
    summarize_text

    verify once again if the revised tools adheres to user feedback. dont make any decision of your own, just stick to user's feedback

    Revised Tools:
    """

        revised_tools_str = self.llm_call(prompt, llm_type)
        return revised_tools_str
    
    # In LLMManager class


# --- Tool Executor ---
class ToolExecutor:
    def __init__(self, tools_config_file, llm_manager):
        self.tools = self._load_tools(tools_config_file)
        self.llm_manager = llm_manager


    def _load_tools(self, tools_config_file: str, selected_agents: List[str] = None) -> Dict:
        """Loads tools from the configuration file, filtered by selected agents."""
        try:
            with open(tools_config_file, "r") as f:
                config = json.load(f)
        except FileNotFoundError:
            logging.error(f"Tools configuration file not found at {tools_config_file}")
            return {}

        tools = {}
        for tool_name, tool_data in config.items():
            if selected_agents is None:
                try:
                    module_name = tool_data["module"]
                    module = importlib.import_module(f"tools.{module_name}")
                    tool_function = getattr(module, tool_name)
                    tools[tool_name] = {
                        "function": tool_function,
                        "description": tool_data["description"],
                        "memory": tool_data.get("memory", False),
                        "agent": tool_data.get("agent", [])
                    }
                except (ImportError, AttributeError, KeyError) as e:
                    logging.error(f"Error loading tool {tool_name}: {e}")
            elif "agent" not in tool_data:
                try:
                    module_name = tool_data["module"]
                    module = importlib.import_module(f"tools.{module_name}")
                    tool_function = getattr(module, tool_name)
                    tools[tool_name] = {
                        "function": tool_function,
                        "description": tool_data["description"],
                        "memory": tool_data.get("memory", False),
                        "agent": tool_data.get("agent", [])
                    }
                except (ImportError, AttributeError, KeyError) as e:
                    logging.error(f"Error loading tool {tool_name}: {e}")
            elif "All" in tool_data.get("agent", []):
                try:
                    module_name = tool_data["module"]
                    module = importlib.import_module(f"tools.{module_name}")
                    tool_function = getattr(module, tool_name)
                    tools[tool_name] = {
                        "function": tool_function,
                        "description": tool_data["description"],
                        "memory": tool_data.get("memory", False),
                        "agent": tool_data.get("agent", [])
                    }
                except (ImportError, AttributeError, KeyError) as e:
                    logging.error(f"Error loading tool {tool_name}: {e}")
            elif any(agent in tool_data.get("agent", []) for agent in selected_agents):
                try:
                    module_name = tool_data["module"]
                    module = importlib.import_module(f"tools.{module_name}")
                    tool_function = getattr(module, tool_name)
                    tools[tool_name] = {
                        "function": tool_function,
                        "description": tool_data["description"],
                        "memory": tool_data.get("memory", False),
                        "agent": tool_data.get("agent", [])
                    }
                except (ImportError, AttributeError, KeyError) as e:
                    logging.error(f"Error loading tool {tool_name}: {e}")

        return tools

    async def execute_tool(self, tool_name: str, arguments: Dict, memory_manager: Any) -> Any:
        """
        Executes a tool with the given arguments.
        
        Args:
            tool_name: The name of the tool to execute.
            arguments: A dictionary of arguments for the tool.
            memory_manager: The MemoryManager instance for accessing results and memory.

        Returns:
            The result of the tool execution.
        """
        if tool_name not in self.tools:
            raise ValueError(f"Tool '{tool_name}' not found.")

        tool_function = self.tools[tool_name]["function"]

        arguments_for_tool_call = arguments.copy()

        # Add memory_manager and llm_manager to the COPY for the tool function call ONLY
        arguments_for_tool_call["memory_manager"] = memory_manager
        arguments_for_tool_call["llm_manager"] = self.llm_manager

        try:
            result = await tool_function(**arguments_for_tool_call)
            return result
        except Exception as e:
            logging.error(f"Error executing tool {tool_name}: {e}")
            return f"Error executing tool {tool_name}: {e}"



# --- Plan Manager ---
class PlanManager:
    def __init__(self, llm_manager: LLMManager, agent,send_message_func):
        self.llm_manager = llm_manager
        self.agent = agent
        self.send_message = send_message_func
        self.plan_storage: Dict[str, Dict] = {}

    async def generate_plan(self, goal: str, memory_manager: MemoryManager,  selected_tools: List[Dict] = None, llm_type: str = None, *, auto_verifier: bool, selected_agents: List[str]= None, session_id=None) -> Union[List[Dict], Dict]:
        """Generates a plan to achieve a given goal."""

        previous_context_str = json.dumps(memory_manager.get_conversation_history(), indent=2)
        graph = await self.llm_manager._generate_graph_from_goal(goal, self.agent.tool_executor.tools, llm_type, self.agent.instruction_manager.dynamic_example_selector, self.agent.instruction_manager.instruction_retriever, selected_tools,selected_agents,retry_count=0, max_retries=3,previous_context=previous_context_str)

        #if plan_verifier:
            #graph = self.verify_graph(graph, goal, llm_type)

        plan_id = generate_guid()
        self.plan_storage[plan_id] = {"plan_id": plan_id,"graph": graph, "goal": goal, "iteration": 0}
        logging.info(f"Plan stored with ID {plan_id}: {self.plan_storage}")
        #plan_verifier = False

        if auto_verifier:
            max_retries = 3
            retries = 0
            while retries < max_retries:
                llm_client, model_name = self.llm_manager.get_client_and_model(llm_type)
                score, feedback_vg = self.llm_manager._verify_graph_with_llm(goal, graph, self.agent.tool_executor.tools, llm_type, selected_agents)
                if score >= 0.8:  # Use a threshold (e.g., 0.8)
                    logging.info(f"Graph verified successfully by AI (Attempt {retries + 1}).")
                    log_message = f"<b><span style='color: green;'>AI Plan Verifier confirmed the plan is valid (Attempt {retries + 1}) and <b>Correctness Score : </b> {score}.</span></b> <br><b> AI Generated Plan :</b> {graph}"
                    #await self.agent.websocket_queue.put(log_message)
                    # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                    return {"plan_id": plan_id, "graph": graph}
                else:
                    logging.warning(f"Graph verification failed by AI (Attempt {retries + 1}). Feedback: {feedback_vg}.  Graph : {graph} Regenerating graph.")
                    log_message = f"<b><span style='color: orange;'>AI Plan Verifier flagged the plan as invalid (Attempt {retries + 1}).</span></b> <br> <span style='color: orange;'>Feedback: {feedback_vg} and <b>Correctness Score : </b> {score}</span> <br> <b><span style='color: orange;'>Attempting to automatically regenerate the plan...</span></b>"
                    #await self.agent.websocket_queue.put(log_message)
                    # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                    if isinstance(feedback_vg, str): # Feedback is available for regeneration
                        graph = await self.llm_manager._generate_graph_from_goal(
                            goal,
                            self.agent.tool_executor.tools,
                            llm_type,
                            self.agent.instruction_manager.dynamic_example_selector,
                            self.agent.instruction_manager.instruction_retriever,
                            selected_tools,
                            selected_agents,
                            retry_count=0,  # Reset retry count for regeneration
                            max_retries=3,
                            feedback_from_AI_verifier=feedback_vg,  # Pass the feedback!
                            previous_graph=graph
                            )
                        if not graph:
                            logging.error(f"Failed to regenerate the graph after AI verification feedback.")
                            log_message = f"<b><span style='color: red;'>Error: Failed to regenerate the plan after AI Plan Verifier feedback.</span></b>"
                            #await self.agent.websocket_queue.put(log_message)
                            # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                            return {"plan_id": plan_id, "graph": graph, "error": "Failed to regenerate the graph after AI feedback"}
                    else: # No feedback string, but verification failed (unlikely case, but handle it)
                        logging.error(f"Graph validation failed by AI (Attempt {retries + 1}), but no feedback was provided. Regenerating graph without specific feedback.")
                        log_message = f"<b><span style='color: orange;'>AI Plan Verifier flagged the plan as invalid (Attempt {retries + 1}), but no specific feedback provided.</span></b> <br> <b><span style='color: orange;'>Attempting to automatically regenerate the plan without specific feedback...</span></b>"
                        #await self.agent.websocket_queue.put(log_message)
                        # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                        graph = await self.llm_manager._generate_graph_from_goal(
                            goal,
                            self.agent.tool_executor.tools,
                            llm_type,
                            self.agent.instruction_manager.dynamic_example_selector,
                            self.agent.instruction_manager.instruction_retriever,
                            selected_tools,
                            selected_agents,
                            retry_count=0,  # Reset retry count for regeneration
                            max_retries=3,
                            feedback_from_AI_verifier="The previous plan was invalid and needs to be revised.Please think step by step and regenerate",
                            previous_graph=graph # Pass the feedback!
                        )
                        if not graph:
                            logging.error(f"Failed to regenerate the graph even without specific AI verification feedback.")
                            log_message = f"<b><span style='color: red;'>Error: Failed to regenerate the plan even without specific AI Plan Verifier feedback.</span></b>"
                            #await self.agent.websocket_queue.put(log_message)
                            # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                            return {"plan_id": plan_id, "graph": graph, "error": "Failed to regenerate the graph after AI feedback (even generic feedback)."}
                    if graph: # Check if graph regeneration was successful
                        logging.info(f"Successfully regenerated graph after AI verification feedback (Attempt {retries + 1}). Retrying verification.")
                        log_message = f"<b><span style='color: green;'>Successfully regenerated plan after AI Plan Verifier feedback (Attempt {retries + 1}).</span></b> <br> <b><span style='color: green;'>Retrying plan verification with the revised plan...</span></b>"
                        #await self.agent.websocket_queue.put(log_message)
                        # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                        #update the plan in plan storage with the revised one.
                        plan_data = self.plan_storage[plan_id]
                        if plan_id:  # <---- Use plan_id_to_update to update existing entry
                            self.plan_storage[plan_id]['graph'] = graph  # <---- UPDATE GRAPH FOR EXISTING PLAN_ID
                            self.plan_storage[plan_id]['iteration'] += 1  # Increment iteration count

                    else: # Graph regeneration failed
                        logging.error(f"Graph regeneration failed after AI verification feedback (Attempt {retries + 1}). Auto-verification failed.")
                        log_message = f"<b><span style='color: red;'>Error: Graph regeneration failed after AI Plan Verifier feedback (Attempt {retries + 1}).</span></b> <br> <span style='color: red;'>Auto-plan verification and correction process failed.</span>"
                        #await self.agent.websocket_queue.put(log_message)
                        #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                        return {"plan_id": plan_id, "graph": graph, "error": "Graph regeneration failed after AI verification feedback. Auto-verification failed."}

                    retries += 1 # Increment retry count for graph regeneration

            logging.error(f"Graph verification failed by AI after multiple attempts. Auto-verification and correction failed.")
            log_message = f"<b><span style='color: red;'>Error: Graph verification failed by AI after multiple attempts ({max_retries}).</span></b> <br> <span style='color: red;'>Auto-plan verification and correction process failed after maximum retries.</span>"
            #await self.agent.websocket_queue.put(log_message)
            #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
            return {"plan_id": plan_id, "graph": graph, "error": f"Graph verification failed by AI after multiple attempts. Auto-verification and correction failed after maximum retries ({max_retries})."}

        return {"plan_id": plan_id, "graph": graph}


    def _validate_graph(self, graph, tools):
        # Placeholder for basic graph validation logic
        # Placeholder for basic graph validation logic
        if not graph:
            return "Graph is empty."

        if "start" not in graph:
            return "Graph is missing a 'start' node."

        if "end" not in graph:
            return "Graph is missing an 'end' node."

        for node_name, node_data in graph.items():
            if node_name in ["start", "end"]:
                continue
            
            if not isinstance(node_data, dict):
                return f"Node '{node_name}' data is not a dictionary."

            if "tool" not in node_data:
                return f"Node '{node_name}' is missing a 'tool' specification."

            tool_name = node_data["tool"]
            if tool_name not in tools:
                return f"Tool '{tool_name}' specified in node '{node_name}' is not a valid tool."
            
            # Check for 'input' key in each node other than 'start' and 'end'
            if "input" not in node_data:
                return f"Node '{node_name}' is missing an 'input' specification."

            if "next" in node_data:
              next_nodes = node_data["next"]
              if isinstance(next_nodes, str):
                  next_nodes = [next_nodes]  # Convert string to list
              
              if not isinstance(next_nodes, list):
                  return f"'next' field in node '{node_name}' should be a list or string."
              
              for next_node in next_nodes:
                if next_node not in graph:
                    return f"Node '{node_name}' references a non-existent next node '{next_node}'."

        # If all checks pass
        return True

    async def execute_plan(self, request: RunAgentRequest, plan: Dict, memory_manager: MemoryManager, tool_executor: ToolExecutor, argument_manager: Any, session_id: str, llm_type: str = None, arguments_verifier: bool = False, get_user_input_func=None, retries=0, max_retries=3) -> Any:
        plan_id = plan.get("plan_id")
        print(f"execute_plan called with plan_id: {plan_id}")
        if plan_id not in self.plan_storage:
            logging.error(f"Plan with ID '{plan_id}' not found.")
            raise ValueError(f"Plan with ID '{plan_id}' not found.")

        plan_data = self.plan_storage[plan_id]
        graph = plan_data.get("graph")
        #memory = {}
        #results = {}
        result = await self._execute_plan_from_graph(request,tool_executor.tools, memory_manager,argument_manager, self.llm_manager.get_client_and_model(llm_type)[0], self.llm_manager.get_client_and_model(llm_type)[1], session_id, llm_type=llm_type, arguments_verifier=arguments_verifier, get_user_input_func=get_user_input_func,retries=0,max_retries=3,plan_id=plan_id)
        return result
    
    
    async def _execute_plan_from_graph(self, request: RunAgentRequest, tools: Dict, memory_manager: MemoryManager, argument_manager: Any, client: Union[AzureOpenAI, genai.GenerativeModel], model_deployment_name: str, session_id: str, llm_type: str = None, arguments_verifier: bool = False, get_user_input_func=None, plan_id=None, current_node: str = "start", retries=0, max_retries=3,plan_regeneration_retries=0, max_plan_regeneration_retries=3) -> Any:

        """
        Executes the plan based on the generated graph.
        _execute_plan_from_graph is designed for the execution of the plan, following the paths determined by the next attributes of the nodes.
        Eg:
        {
        "start": ["A"],
        "A": {"next": ["B", "C"]},
        "B": {"next": ["D"]},
        "C": {"next": ["E"]},
        "D": {"next": ["end"]},
        "E": {"next": ["end"]},
        "end": {}
        }
        _execute_plan_from_graph Order (ignoring the retries and other exception logic): The execution order will be like A -> B -> D-> end and then A -> C -> E -> end. The entire plan execution will take two iterations.
        """
        logging.info(f"Debugging: _execute_plan_from_graph called with plan_id: {plan_id}")

        """
        # Get the original goal from plan_storage
        original_goal = ""
        #plan_id = None # Initialize plan_id
        for pid, plan_data in self.plan_storage.items():
            print("entering pid in exec plan---plan_data[graph]", plan_data['graph'])
            print("entering pid in exec plan---graph", graph)
            if plan_data['graph'] == graph:
                original_goal = plan_data['goal']
                plan_id = pid
                break
        if not plan_id:
            logging.error(f"No plan with graph found in plan_storage.")
            return None  # Exit if plan_id is not found

        # Fetch the latest graph from plan_storage
        plan_data = self.plan_storage[plan_id]
        graph = plan_data.get("graph") """

        if plan_id is None: # <---- Check if plan_id is None (should not happen in normal flow after correction in PlanManager.execute_plan)
            logging.error(f"Error: plan_id is None when _execute_plan_from_graph is called.") # Add error log
            return None

        if plan_id not in self.plan_storage: # Check plan_storage using plan_id
            logging.error(f"No plan with ID '{plan_id}' found in plan_storage.")
            return None

        plan_data = self.plan_storage[plan_id] # Retrieve plan_data using plan_id (directly and efficiently)
        original_goal = plan_data['goal']
        graph = plan_data.get("graph")

        logging.info(f"Executing plan from graph: {graph}")
        #log_message = json.dumps({"type": MessageType.LOG.value,"message": self.format_plan_execution_message()})
        #log_message = json.dumps({"type": MessageType.LOG.value, "message": "<b><span style='color: black;'>Executing plan from graph...</span></b>"})
        #log_message = self.format_plan_execution_message()
        #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
        #await self.agent.websocket_queue.put(log_message)
        await asyncio.sleep(0.1)
        #await websocket_queues[session_id].put(log_message)

        

        # Get the nodes from start node for sequential execution
        if current_node == "start":
            current_nodes = graph.get("start",[])
        else:
            current_nodes = [current_node]
        logging.info(f"Starting Nodes:{current_nodes}")

        # Get the nodes based on the topological order
        #ordered_nodes = _topological_sort(graph)
        #logging.info(f"Topological ordered nodes: {ordered_nodes}")

        # Start executing from the current_node
        executed_nodes = set() # Add a set to store executed nodes
        #verified_nodes = set()

        failed_node_counts = {} 
        results = {}  # Initialize results here!
        
        while current_nodes:
            next_nodes = []
            tasks = [] # Initialize tasks list for parallel execution
            for node in current_nodes:
                if node == "start" or node == "end" or node in executed_nodes:
                   continue

                # Check if the node is frozen.  If so, skip it *entirely*.
                """if graph[node].get("frozen", False) == True: #if frozen, dont execute
                    logging.info(f"Skipping frozen node: {node}")
                    # Still need to add next nodes, as usual.
                    if graph[node].get("next"):
                        next_nodes.extend(graph[node].get("next", []))
                    executed_nodes.add(node)
                    continue """

            # --------------------------------------

                executed_nodes.add(node)
               
                async def execute_single_node(node_id ,graph, tools, memory_manager, argument_manager): # Define inner async function for task
                    tool_name = None # Initialize tool_name here
                    arguments = None # Initialize arguments here
                    node = node_id

                    try:               
                        print(f"Current node in loop: {node}")

                        tool_name, arguments = argument_manager._determine_tool_and_arguments_for_node(node, graph, memory_manager, tools, client, model_deployment_name, llm_type,retry_count=0, max_retries=3)
                        tool_name = graph[node].get("tool")

                        if not tool_name:
                            logging.warning(f"Tool name not found in node: '{node}'. Skipping node execution.")
                            return node_id, None 

                        if tool_name not in tools:
                            logging.error(f"Tool '{tool_name}' not found.")
                            raise ValueError(f"Tool '{tool_name}' not found.")

                        arguments = memory_manager._substitute_placeholders(arguments)
                        #arguments = await memory_manager._llm_substitute_placeholders(arguments, results, self.llm_manager) # Pass llm_manager

                        #if arguments_verifier and node not in verified_nodes:
                        if request.arguments_verifier:
                                # User verification
                                if not await self.agent.verify_arguments(tool_name, arguments, node, graph,memory_manager, llm_type, get_user_input_func, session_id):
                                    logging.error(f"Argument verification failed for node '{node}'. Exiting.")
                                    return
                        else:
                                await asyncio.sleep(0.1)
                                print(f"Skipping argument verification for tool '{tool_name}' at node '{node}' as arguments_verifier is False.")

                        logging.info(f"Executing tool: {tool_name} with arguments: {arguments}")

                        result = await self.agent.tool_executor.execute_tool(tool_name, arguments, memory_manager)
                        print( "tool verifier flag", request.tool_result_verifier)
                        ai_feedback = None
                        if request.tool_result_verifier: # Check the new flag
                                feedback_history = [] # <---- Initialize feedback history list
                                tool_verification_result = self.llm_manager._verify_tool_result_with_llm(
                                    tool_name,
                                    self.agent.tool_executor.tools[tool_name]["description"],
                                    arguments,
                                    result,
                                    original_goal, # Assuming original_goal is accessible here
                                    memory_manager,
                                    llm_type
                                )
                                if tool_verification_result != True:
                                    ai_feedback = tool_verification_result
                                    logging.warning(f"Tool result verification failed for node '{node}'. Feedback: {tool_verification_result}")
                                    feedback_history.append(tool_verification_result) 
                                    # --- Automatic Argument Regeneration and Retry ---
                                    if retries < max_retries: # Check retry count
                                        logging.info(f"Attempting to regenerate arguments for node '{node}' (Retry {retries + 1}/{max_retries}).")
                                        #log_message = json.dumps({"type": MessageType.LOG.value, "message": f"<b><span style='color: orange;'>AI Tool Result Verifier flagged the result of '{tool_name}' as potentially invalid.</span></b> <br> <span style='color: orange;'>Feedback: {tool_verification_result}</span> <br> <b><span style='color: orange;'>Attempting to automatically regenerate arguments and retry...</span></b>"})
                                        #log_message = self.format_verifier_message(tool_name, tool_verification_result)
                                        #await self.agent.websocket_queue.put(log_message)
                                        #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                                        # Regenerate arguments using feedback from tool result verifier
                                        new_arguments = await argument_manager._regenerate_arguments_with_AI_feedback(
                                            tool_name,
                                            self.agent.tool_executor.tools[tool_name]["description"],
                                            graph,
                                            memory_manager,
                                            tool_verification_result, # Pass the verifier's feedback
                                            node,
                                            llm_type
                                        )

                                        if new_arguments:
                                            logging.info(f"Successfully regenerated arguments for tool '{tool_name}' at node '{node}'. Retrying execution.")
                                            # Update the graph with new arguments (similar to how you do in manual feedback loop)
                                            for node_item in graph:
                                                if node_item == node:
                                                    if isinstance(graph[node], dict):
                                                        if "input" in graph[node]:
                                                            graph[node]["input"].clear()
                                                            graph[node]["input"].update(new_arguments)
                                                        else:
                                                            graph[node]["input"] = new_arguments
                                                    break

                                            # Recursively retry execution from the current node with incremented retry count
                                            return await self._execute_plan_from_graph(
                                                request=request, # Pass request object
                                                tools=tools,
                                                #graph=graph,
                                                memory_manager=memory_manager,
                                                argument_manager=argument_manager,
                                                client=client,
                                                model_deployment_name=model_deployment_name,
                                                session_id=session_id,
                                                llm_type=llm_type,
                                                arguments_verifier=arguments_verifier,
                                                get_user_input_func=get_user_input_func,
                                                current_node=node, # Retry from the same node
                                                retries=retries + 1, # Increment retry count
                                                max_retries=max_retries,
                                                plan_id=plan_id,
                                                plan_regeneration_retries=plan_regeneration_retries, max_plan_regeneration_retries=max_plan_regeneration_retries
                                            )
                                        else:
                                            logging.error(f"Failed to regenerate arguments for tool '{tool_name}' at node '{node}' even after tool result verification failure feedback. Proceeding without retry.")
                                            log_message =  f"<b><span style='color: red;'>Error: Failed to regenerate arguments for '{tool_name}' after AI Tool Result Verifier flagged it as invalid.</span></b> <br> <span style='color: orange;'>Proceeding without automatic retry. Result might be unreliable.</span>"
                                            #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                                            return node, None 


                                    else: # Max retries exceeded for ARGUMENT regeneration, now handle PLAN

                                        logging.warning(f"Maximum retries ({max_retries}) exceeded for node '{node}' after tool result verification failure. Proceeding without further retries.")
                                        if plan_regeneration_retries < max_plan_regeneration_retries: 
                                            logging.warning(f"Persistent tool RESULT failure detected for node '{node}'. Triggering PLAN regeneration (Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries}).")  # <--- Changed log message

                                            log_message = f"<b><span style='color: red;'>Persistent tool RESULT failure detected for '{tool_name}' at node '{node}'.</span></b> <br> <b><span style='color: red;'>Attempting to regenerate the PLAN (Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries})...</span></b>"

                                            #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                                            selected_agents = request.agents

                                            # Trigger plan regeneration with feedback about persistent tool failure
                                            revised_graph = await self._regenerate_graph_with_feedback_for_persistent_failure(
                                                goal=original_goal,
                                                failed_tool_node=node,
                                                failed_tool_name=tool_name,
                                                previous_graph=graph,
                                                memory_manager = memory_manager,
                                                llm_type=llm_type,
                                                selected_agents=selected_agents,
                                                feedback_from_retries="\n".join(feedback_history) # <---- Pass feedback here
                                            )

                                            if revised_graph:
                                                logging.info(f"Successfully regenerated PLAN after persistent tool result failure (Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries}).")

                                                self.plan_storage[plan_id]['graph'] = revised_graph
                                                self.plan_storage[plan_id]['iteration'] += 1
                                                logging.info(f"Updated plan_storage for plan_id {plan_id} with new graph: {revised_graph}")

                                            
                                                    # Update the plan in plan_storage with the revised graph
                                                #plan_id_to_update = plan_data.get("plan_id") 
                                                #if plan_id_to_update: # <---- Use plan_id_to_update to update existing entry
                                                 #   self.plan_storage[plan_id_to_update]['graph'] = revised_graph # <---- UPDATE GRAPH FOR EXISTING PLAN_ID
                                                  #  self.plan_storage[plan_id_to_update]['iteration'] += 1 # Increment iteration count
                                                #else:
                                                 #   logging.error("Error: Could not retrieve plan_id to update plan storage with revised graph.") # Handle error if plan_id is missing

                                                # Recursively retry plan execution from the start with the revised graph

                                                logging.info(f"Retrying plan execution with revised graph (Regeneration Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries}).")

                                                return await self._execute_plan_from_graph( # Recursive call with revised graph
                                                    request=request, # Pass request
                                                    tools=tools,
                                                    #graph=revised_graph, # Use revised_graph
                                                    plan_id=plan_id,
                                                    memory_manager=memory_manager,
                                                    argument_manager=argument_manager,
                                                    client=client,
                                                    model_deployment_name=model_deployment_name,
                                                    session_id=session_id,
                                                    llm_type=llm_type,
                                                    arguments_verifier=arguments_verifier,
                                                    get_user_input_func=get_user_input_func,
                                                    retries=0, # Reset retries for the new plan
                                                    max_retries=max_retries,
                                                    current_node = "start"
                                                )
                                            else:
                                                logging.error(f"Failed to regenerate PLAN after persistent tool result failure for node '{node}' (Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries}). Proceeding with potentially flawed plan.")
                                                
                                                log_message = f"<b><span style='color: red;'>Error: Failed to regenerate PLAN even after persistent tool result failure for '{tool_name}' at node '{node}' (Attempt {plan_regeneration_retries + 1}/{max_plan_regeneration_retries}).</span></b> <br> <span style='color: orange;'>Proceeding with potentially flawed plan.</span>"

                                                #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                                        else: # Max PLAN regeneration retries exceeded!
                                            logging.error(f"Maximum PLAN regeneration retries ({max_plan_regeneration_retries}) exceeded for node '{node}'. Proceeding with potentially flawed plan.") # <--- New Error Log
                                            
                                            # log_message = f"<b><span style='color: red;'>Error: Maximum PLAN regeneration retries exceeded for '{tool_name}' at node '{node}' ({max_plan_regeneration_retries}).</span></b> <br> <span style='color: orange;'>Proceeding with potentially flawed plan.</span>" # <--- New Error Log
                                            #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                                    # Continue execution even if verification failed or retries failed (for now - you can customize this)
                                    logging.warning(f"Tool result deemed invalid by AI: {tool_verification_result}. Proceeding but result might be unreliable.")
                                    log_message = f"<b><span style='color: orange;'>Warning: AI Tool Result Verifier flagged the result of '{tool_name}' as potentially invalid:</span></b> <br> <span style='color: orange;'>Feedback: {tool_verification_result}</span> <br> <span style='color: orange;'>Proceeding but result might be unreliable.</span>"
                                    #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                                else:
                                    logging.info(f"Tool result verified successfully by AI.")
                                    #log_message = json.dumps({"type": MessageType.LOG.value, "message": f"<b><span style='color: blue;'>AI Tool Result Verifier confirmed the result of '{tool_name}' is valid.</span></b>"})
                                    #log_message = json.dumps({"type": MessageType.LOG.value, "message": f"<b><span style='color: blue;'>AI Tool Result Verifier confirmed the result of '{tool_name}' is valid.</span></b>"})
                                    log_message = f"<div style='background-color: #85EAFC; padding: 8px 12px; border-radius: 4px; border-left: 4px solid #2196F3;'>"
                                    log_message+= f"<span style='display: inline-flex; align-items: center; gap: 8px;'>"
                                    log_message+= f"<span style='color: #2196F3; font-weight: bold;'>✓</span>"
                                    log_message+= f"<span style='color: #0d47a1;'>AI Tool Result Verifier confirmed the result of '{tool_name}' is valid.</span>"
                                    log_message+= f"</span>"
                                    log_message+= f"</div>"
                                    #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                        request.human_tool_result_verifier = True
                        #if request.human_tool_result_verifier:
                        if request.arguments_verifier:
                            #prompt = self.agent.format_user_input_prompt(arguments, tool_name=tool_name)
                            prompt = f"<p><b>Tool Output:</b><br><pre>{json.dumps(result, indent=2)}</pre></p>"
                            if ai_feedback:
                                prompt += f"<p><b>AI Feedback:</b> {ai_feedback}</p>"
                            prompt += "<p>Is this result correct? Type 'yes' to continue, or provide feedback to revise.</p>"
                            user_feedback = await get_user_input_func(f"{prompt}", session_id=session_id)
                            if user_feedback.lower() != "yes":
                                logging.info(f"Feedback received: {user_feedback}. Revising the arguments...")
                                memory_manager.results[f"user_feedback_{node}"] = user_feedback
                                new_arguments = await argument_manager._regenerate_arguments_with_feedback(
                                    tool_name,
                                    self.agent.tool_executor.tools[tool_name]["description"],
                                    graph,
                                    memory_manager,
                                    user_feedback,
                                    node,
                                    llm_type
                                )
                                if new_arguments:
                                    for node_item in graph:
                                        if node_item == node:
                                            if isinstance(graph[node], dict):
                                                if "input" in graph[node]:
                                                    graph[node]["input"].clear()
                                                    graph[node]["input"].update(new_arguments)
                                                else:
                                                    graph[node]["input"] = new_arguments
                                            break
                                    return await self._execute_plan_from_graph(
                                        request=request,
                                        tools=tools,
                                        memory_manager=memory_manager,
                                        argument_manager=argument_manager,
                                        client=client,
                                        model_deployment_name=model_deployment_name,
                                        session_id=session_id,
                                        current_node=node,
                                        llm_type=llm_type,
                                        arguments_verifier=arguments_verifier,
                                        get_user_input_func=get_user_input_func,
                                        retries=retries + 1,
                                        max_retries=max_retries,
                                        plan_id=plan_id
                                    )
                                else:
                                    logging.error(f"Failed to revise arguments for tool '{tool_name}' at node '{node}'.")
                                    return node, None
                            else:
                                logging.info(f"Tool result verified successfully by user at node '{node}'.")
                                log_message = f"<div style='background-color: #85EAFC; padding: 8px 12px; border-radius: 4px; border-left: 4px solid #2196F3;'>"
                                log_message += f"<span style='display: inline-flex; align-items: center; gap: 8px;'>"
                                log_message += f"<span style='color: #2196F3; font-weight: bold;'>✓</span>"
                                log_message += f"<span style='color: #0d47a1;'>Tool Result verified by User for tool '{tool_name}'.</span>"
                                log_message += f"</span></div>"
                                #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                        print(f"  --> Result: {result}")
                        logging.info(f"Tool '{tool_name}' executed successfully. Result: {result}")
                        await asyncio.sleep(0.1)

                        # Create a serializable copy of the arguments
                        serializable_arguments = {k: v for k, v in arguments.items() if k not in ["memory_manager", "llm_manager"]}

                        result_entry = {
                                    "node_id": node,
                                    "tool": tool_name,
                                    "input": serializable_arguments,
                                    "result": result
                        }
                        memory_manager.store_result(node, result_entry)
                        #graph[node]['frozen'] = True # set the frozen attribute

                        print(f"  --> results after executing {tool_name}: {memory_manager.results}")  # Print the contents of 'results'

                        if tools[tool_name].get("memory", False):
                            memory_manager.store_in_memory(node, result)

                        # Fetch the latest graph from plan_storage after each node execution
                        #plan_data = self.plan_storage[plan_id]
                        #graph = plan_data.get("graph")
                        if graph[node].get("next"):
                            next_nodes.extend(graph[node].get("next",[]))
                        #executed_nodes.add(node) # Add node to executed set
                        
                        return node_id, result_entry # Return node_id and result_entry for 
                    
                    except Exception as e:
                        logging.error(f"Error during execution of node '{node_id}': {e}")
                        if tool_name is not None and arguments is not None:  # Check if tool_name and arguments were defined
                            return node_id, (e, tool_name, arguments) # Return exception, tool_name, and arguments
                        else:
                            return node_id, (e, tool_name, arguments)
                
                task = asyncio.create_task(execute_single_node(node,graph,tools,memory_manager,argument_manager)) # Create task for each node
                tasks.append(task) # Add task to task list
            
            results_concurrent = await asyncio.gather(*tasks, return_exceptions=True) # Run tasks in parallel

            for result_item in results_concurrent: # Process results from concurrent tasks
               if isinstance(result_item, HTTPException): # Check if result is HTTPException (argument verifier fail)
                    return result_item # Return HTTPException to stop execution
               if isinstance(result_item, tuple) and len(result_item) == 2: # Check if result is tuple (node_id, result or exception)
                    node_id, execution_result = result_item
                    if isinstance(execution_result, dict):
                    # Success case: execution_result is result_entry
                        memory_manager.store_result(node_id, execution_result)
                        if graph[node_id].get("next"):
                            next_nodes.extend(graph[node_id].get("next", []))
                        executed_nodes.add(node_id)

                    elif isinstance(execution_result, Exception): # Check if execution resulted in exception
                        e, tool_name, arguments = execution_result
                        logging.error(f"Error during execution of node '{node_id}': {execution_result}")
                        if isinstance(execution_result, ValueError):
                            if "Tool name extraction failed" in str(execution_result):
                                continue  # Skip if the tool name failed to be extracted

                        #check if tool_name and arguments exists to avoid error in feedback prompt

                        if 'tool_name' in locals() and 'arguments' in locals():
                            tool_name = graph[node_id].get("tool", "Unknown Tool") # Fallback tool_name - use node_id or "Unknown Tool"
                            logging.warning(f"tool_name was not properly defined due to exception. Using fallback tool name: '{tool_name}'.") # Log fallback

                            failed_node_counts[node] = failed_node_counts.get(node, 0) + 1 # <----- ADD THIS LINE: Increment failure count

                            if retries >= max_retries:
                                logging.info(f"Maximum retries exceeded for node '{node}'. Asking user for feedback.")

                            # --- ADD the Plan Regeneration Trigger Logic Here --- <----------------- ADD BLOCK
                                if failed_node_counts.get(node, 0) >= max_retries: # Check failure count
                                        logging.warning(f"Persistent tool failure detected for node '{node}'. Triggering plan regeneration.")
                                        log_message = f"<b><span style='color: red;'>Persistent tool failure detected for '{tool_name}' at node '{node}'.</span></b> <br> <b><span style='color: red;'>Attempting to regenerate the plan...</span></b>"
                                        #await self.agent.websocket_queue.put(log_message)
                                        #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)

                                        selected_agents = request.agents 

                                        # Trigger plan regeneration with feedback about persistent tool failure
                                        revised_graph = await self._regenerate_graph_with_feedback_for_persistent_failure(
                                            goal=original_goal,
                                            failed_tool_node=node,
                                            failed_tool_name=tool_name,
                                            previous_graph=graph,
                                            memory_manager = memory_manager,
                                            llm_type=llm_type,
                                            selected_agents=selected_agents
                                        )

                                        if revised_graph:
                                            logging.info(f"Successfully regenerated plan after persistent tool failure.")
                                            # Update the plan in plan_storage with the revised graph
                                            #plan_id = plan_data.get("plan_id") # Assuming plan_data is still accessible here
                                            plan_id = [pid for pid, p_data in self.plan_storage.items() if p_data['graph'] == graph]
                                            if plan_id:
                                                plan_id = plan_id[0] #plan id is returned as list, so reading the first element.
                                                self.plan_storage[plan_id]['graph'] = revised_graph
                                                self.plan_storage[plan_id]['iteration'] += 1 # Increment iteration count

                                            # Recursively retry plan execution from the start with the revised graph
                                            logging.info(f"Retrying plan execution with revised graph.")
                                            return await self._execute_plan_from_graph( # Recursive call with revised graph
                                                request=request, # Pass request
                                                tools=tools,
                                                #graph=revised_graph, # Use revised_graph
                                                memory_manager=memory_manager,
                                                argument_manager=argument_manager,
                                                client=client,
                                                model_deployment_name=model_deployment_name,
                                                session_id=session_id,
                                                llm_type=llm_type,
                                                arguments_verifier=arguments_verifier,
                                                get_user_input_func=get_user_input_func,
                                                retries=0, # Reset retries for the new plan
                                                max_retries=max_retries,
                                                plan_regeneration_retries = plan_regeneration_retries+1,
                                                current_node = "start"
                                            )
                                        else:
                                            logging.error(f"Failed to regenerate plan after persistent tool failure for node '{node}'. Proceeding with potentially flawed plan.")
                                            log_message = f"<b><span style='color: red;'>Error: Failed to regenerate plan even after persistent tool failure for '{tool_name}' at node '{node}'.</span></b> <br> <span style='color: orange;'>Proceeding with potentially flawed plan.</span>"
                                            #await self.agent.websocket_queue.put(log_message)
                                            #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
                                else: # if failure count is less than max_retries, proceed with existing logic
                                    # Get user feedback
                                    #feedback = input(f"Error during execution of node '{node}'. Provide feedback to revise the arguments or type 'skip' to skip this node: ")
                                    feedback = await get_user_input_func(f"Error during execution of node '{node}'. Provide feedback to revise the arguments or type 'skip' to skip this node: ",session_id=session_id)
                                    if feedback.lower() == 'skip':
                                            continue  # Skip the current node and proceed with the next one
                                        # Use the user feedback to regenerate arguments
                                    llm_client, model_name = self.llm_manager.get_client_and_model(llm_type)
                                    new_arguments = await argument_manager._regenerate_arguments_with_feedback(
                                            tool_name,
                                            self.agent.tool_executor.tools[tool_name]["description"],
                                            graph,
                                            memory_manager,
                                            feedback,  # Use user-provided feedback
                                            node,
                                            llm_type)
                            else:
                                # Use LLM to generate feedback prompt
                                llm_client, model_name = self.llm_manager.get_client_and_model(llm_type)
                                tool_descriptions = _get_tool_descriptions(tools)
                                feedback_prompt = self.llm_manager._error_to_feedback_prompt(e, tool_name, arguments, original_goal, memory_manager, tool_descriptions)

                                    # Regenerate arguments with feedback
                                new_arguments = await argument_manager._regenerate_arguments_with_feedback(
                                        tool_name,
                                        self.agent.tool_executor.tools[tool_name]["description"],                                
                                        graph,
                                        memory_manager,
                                        feedback_prompt,
                                        node,
                                        llm_type
                                    )

                            if new_arguments is None:
                                logging.error(f"Failed to regenerate arguments for node '{node}'. Exiting plan execution.")
                                return  # Exit if argument regeneration fails

                            # Update the graph with new arguments
                            for node_item in graph:
                                if node_item == node:
                                    if isinstance(graph[node], dict):
                                        if "input" in graph[node]:
                                            graph[node]["input"].clear()
                                            graph[node]["input"].update(new_arguments)
                                        else:
                                            graph[node]["input"] = new_arguments
                                    else:
                                        arguments.clear()
                                        arguments.update(new_arguments)
                                    break

                            # Find the plan in plan_storage and update the graph
                            """for plan_id, plan_data in self.plan_storage.items():
                                if plan_data['graph'] == graph:
                                    self.plan_storage[plan_id]['graph'] = graph
                                    break"""

                            # Recursively retry execution from the current node
                            logging.info(f"Retrying execution from node '{node}' with new arguments.")

                            
                            #return await self._execute_plan_from_graph(tools, graph, memory_manager, argument_manager,client, model_deployment_name, session_id, node, llm_type, arguments_verifier, get_user_input_func, retries + 1, max_retries)
                            return await self._execute_plan_from_graph(
                                request=request, # Pass request
                                tools=tools,
                                #graph=graph, # Pass the GRAPH correctly
                                memory_manager=memory_manager, # Pass memory_manager correctly
                                argument_manager=argument_manager,
                                client=client,
                                model_deployment_name=model_deployment_name,
                                session_id=session_id,
                                current_node=node,  # Pass current_node - IMPORTANT: make sure you are passing current_node here, not 'node' which might be undefined outside the loop
                                llm_type=llm_type,
                                arguments_verifier=arguments_verifier,
                                get_user_input_func=get_user_input_func,
                                retries=retries + 1,
                                max_retries=max_retries,
                                plan_id = plan_id
                            )
                    else:
                        logging.warning(f"Unexpected execution_result type: {type(execution_result)}")
                        #executed_nodes.add(node_id) # Add node to executed set
               elif result_item is None: # Handle skipped node (tool_name not found)
                    continue # Skip to next result_item
               else: # Handle unexpected result type (optional - for robustness)
                    logging.warning(f"Unexpected result type from task execution: {type(result_item)}, value: {result_item}")
                
            current_nodes = next_nodes # set the next node
        return {"steps": list(memory_manager.results.values())}

    
    # In PlanManager class:

    async def _regenerate_graph_with_feedback_for_persistent_failure(self, goal: str, failed_tool_node: str, failed_tool_name: str, previous_graph: Dict, memory_manager: MemoryManager,llm_type: str = None,selected_agents: List[str] = None, feedback_from_retries: str = None) -> Dict: # <--- Added feedback_from_retries
        """
        Helper function to regenerate the graph specifically for persistent tool failures.
        """
        feedback_plan = f"The tool '{failed_tool_name}' at node '{failed_tool_node}' has consistently failed even after multiple attempts to regenerate arguments. The current plan is likely flawed and needs to be revised to avoid using this tool or find an alternative approach to achieve the goal."

        if feedback_from_retries: # <--- Check if feedback is available
            feedback_plan += f"\n\n**Detailed Feedback from Tool Result Verifier (during retries):**\n{feedback_from_retries}" # <--- Append to feedback_plan

        revised_graph = await self.llm_manager._regenerate_graph_with_AI_feedback(
            goal, self.agent.tool_executor.tools, feedback_plan, previous_graph,memory_manager,llm_type,selected_agents=selected_agents
        )
        return revised_graph
    

# --- Argument Manager ---
class ArgumentManager:
    def __init__(self, llm_manager: LLMManager, memory_manager: MemoryManager,send_message_func):
        self.llm_manager = llm_manager
        self.memory_manager = memory_manager # Store memory_manager as self.memory_manager
        self.send_message = send_message_func

    def _determine_tool_and_arguments_for_node_prev(self, node: str, graph: Dict, memory_manager: MemoryManager,tools: Dict, client: Union[AzureOpenAI, genai.GenerativeModel], model_deployment_name: str, llm_type: str = None) -> tuple[str, Dict]:
        """
        Determines the tool name and arguments for a given node in the graph using the LLM.
        Now handles the new graph format where entities are stored as attributes.
        """

        print(f"\n--- Results at start of _determine_tool_and_arguments_for_node (Node: {node}) ---")
        print(json.dumps(memory_manager.results, indent=2))
        print("--- End Results ---\n")

        print("\n--- Printing self.memory_manager.results in determine tool ---")
        print(json.dumps(memory_manager.results, indent=2)) 
        print("--- End ---")
        
        # Extract tool name and arguments from the node in the graph
        tool_name = graph[node].get("tool")

        # Get the tool description
        tool_description = tools.get(tool_name, {}).get("description", "")
        """matching_tool_name = None
        matching_tool_description = None
        for key, tool_data in tools.items():
            if key.lower() == tool_name.lower(): # Case-insensitive comparison
                matching_tool_name = key
                matching_tool_description = tool_data.get("description", "")
                break # Found a match, exit loop

        if matching_tool_name:
            tool_name = matching_tool_name # Use the actual key from tools.json for consistency
            tool_description = matching_tool_description
        else:
            tool_name = None # Set to None if no match found (even case-insensitive)
            tool_description = None"""


        if not tool_name or not tool_description:
            logging.error(f"Tool name extraction failed or tool '{tool_name}' not found or has no description in tools.json.")
            raise ValueError(f"Tool name extraction failed or tool '{tool_name}' not found or has no description in tools.json.")
        
        # Extract arguments from node attributes
        arguments = graph[node].get("input", {})

        # Use the LLM to generate arguments
        prompt = f"""
    You are an AI assistant that determines the arguments for a tool based on the current node in a graph, the overall goal, the tool's description, and the results of previous steps.

    **User Goal:** {memory_manager.get_from_memory("user_goal")}  
    **Current Node:** {node}
    **Graph:** {json.dumps(graph)}
    **Tool Name:** {tool_name}
    **Tool Description:** {tool_description}
    **Current Memory (Results of Previous Steps):** {json.dumps(memory_manager.results)}

    **Task:**
    Generate the arguments for the tool in JSON format.

    **Important Instructions:**
    - **The output must be a valid JSON string and nothing else. Do not include any introductory or explanatory text.**
    - Use the information from the current node (e.g., entity names, parameters) to determine argument values.
    - Refer to the tool description to understand what arguments are required and their expected format (e.g., string, number, etc.).
    - Use the results of previous steps (available in the `Current Memory`) if the tool needs them as input. Reference previous results using the format `{{node_name}}` (e.g., `{{predict_payment_timing_BOFAMAST}}`).
    - If the goal has a keyword like 'prediction', the email content in 'draft_email' tool MUST include the predicted business day.
    - **If no arguments are needed, return an empty JSON object {{}}.**
    - **For the 'draft_email' tool, ensure that the 'to_recipients' argument is always a JSON array (list) of email address strings. Even if there is only one recipient, it should still be a list containing that single email address.**

    **Examples:**

    Graph: {{"start": ["predict_payment_timing_1", "predict_payment_timing_2"], "predict_payment_timing_1": {{"tool": "predict_payment_timing", "customer": "BOFAMAST", "amount": 50000000, "next": ["draft_email_1"]}}, "predict_payment_timing_2": {{"tool": "predict_payment_timing", "customer": "ABNMST", "amount": 22222, "next": ["draft_email_2"]}}, "draft_email_1": {{"tool": "draft_email", "recipient": "ddd@www.com", "next": ["end"]}}, "draft_email_2": {{"tool": "draft_email", "recipient": "3333@ddd.com", "next": ["end"]}}, "end": {{}}}}
    Current Node: predict_payment_timing_1
    Tool Name: predict_payment_timing
    Tool Description: Predicts the payment timing for a given customer and amount. Arguments: customer_id (required, string), amount (required, number).
    Current Memory: {{}}

    Arguments:
    {{"customer_id": "BOFAMAST", "amount": 50000000}}

    Current Node: draft_email_2
    Tool Name: draft_email
    Tool Description: Drafts an email. Arguments: subject (required, string), recipient (required, string), content (required, string).
    Current Memory: {{"predict_payment_timing_2": "some_prediction_result"}}

    Arguments:
    {{
        "subject": "Payment Prediction",
        "recipient": "3333@ddd.com",
        "content": "The predicted business day for ABNMST is {{predict_payment_timing_2}}"
    }}

    Current Node: summarize_text_1
    Tool Name: summarize_text
    Tool Description: Summarizes the given text. Arguments: text (required, string)
    Current Memory: {{"predict_payment_timing_2": "some_prediction_result"}}

    Arguments:
    {{
        "text": "some_prediction_result"
    }}

    **Generate the arguments for the tool specified by the Current Node:**

    Arguments:
    """
        arguments_str = self.llm_manager.llm_call(prompt, llm_type)

            
        try:
            arguments = json.loads(arguments_str)
            logging.info(f"Prompt for determine node {prompt}.")
        except json.JSONDecodeError as e:
            logging.error(f"Failed to parse JSON response from LLM while generating arguments for node '{node}' due to error {e}.")
            logging.info(f"The invalid JSON string was: {arguments_str}.")
            raise ValueError(
                f"Failed to parse JSON response from LLM while generating arguments for node '{node}' due to error {e}. The invalid JSON string was: {arguments_str}."
            )
        return tool_name, arguments 

    def _determine_tool_and_arguments_for_node(self, node: str, graph: Dict, memory_manager: MemoryManager, tools: Dict, client: Union[AzureOpenAI, genai.GenerativeModel], model_deployment_name: str, llm_type: str = None, retry_count: int = 0, max_retries: int = 3,feedback_prompt_json_error: str = None) -> tuple[str, Dict]: # Add retry parameters
        """
        Determines the tool name and arguments for a given node in the graph using the LLM with retry for JSON errors.
        """
        logging.info(f"--- Determine Arguments for Node: {node} (Attempt {retry_count + 1}) ---") # Log retry attempt

        tool_name = graph[node].get("tool")
        tool_description = tools.get(tool_name, {}).get("description", "")

        if not tool_name or not tool_description:
            logging.error(f"Tool name extraction failed or tool '{tool_name}' not found or has no description in tools.json.")
            raise ValueError(f"Tool name extraction failed or tool '{tool_name}' not found or has no description in tools.json.")

        arguments = graph[node].get("input", {})

        prompt = f"""
    You are an AI assistant that determines the arguments for a tool based on the current node in a graph, the overall goal, the tool's description, and the results of previous steps.

    **User Goal:** {memory_manager.get_from_memory("user_goal")}  
    **Current Node:** {node}
    **Graph:** {json.dumps(graph)}
    **Tool Name:** {tool_name}
    **Tool Description:** {tool_description}
    **Current Memory (Results of Previous Steps):** {json.dumps(memory_manager.results)}

    **Task:**
    Generate the arguments for the tool in JSON format.

    **Important Instructions:**
    - **The output must be a valid JSON string and nothing else. Do not include any introductory or explanatory text.**
    - Use the information from the current node (e.g., entity names, parameters) to determine argument values.
    - Refer to the tool description to understand what arguments are required and their expected format (e.g., string, number, etc.).
    - Use the results of previous steps (available in the `Current Memory`) if the tool needs them as input. Reference previous results using the format `{{node_name}}` (e.g., `{{predict_payment_timing_BOFAMAST}}`).
    - If the goal has a keyword like 'prediction', the email content in 'draft_email' tool MUST include the predicted business day.
    - **If no arguments are needed, return an empty JSON object {{}}.**
    - **For the 'draft_email' tool, ensure that the 'to_recipients' argument is always a JSON array (list) of email address strings. Even if there is only one recipient, it should still be a list containing that single email address.**

    **Examples:**

    Graph: {{"start": ["predict_payment_timing_1", "predict_payment_timing_2"], "predict_payment_timing_1": {{"tool": "predict_payment_timing", "customer": "BOFAMAST", "amount": 50000000, "next": ["draft_email_1"]}}, "predict_payment_timing_2": {{"tool": "predict_payment_timing", "customer": "ABNMST", "amount": 22222, "next": ["draft_email_2"]}}, "draft_email_1": {{"tool": "draft_email", "recipient": "ddd@www.com", "next": ["end"]}}, "draft_email_2": {{"tool": "draft_email", "recipient": "3333@ddd.com", "next": ["end"]}}, "end": {{}}}}
    Current Node: predict_payment_timing_1
    Tool Name: predict_payment_timing
    Tool Description: Predicts the payment timing for a given customer and amount. Arguments: customer_id (required, string), amount (required, number).
    Current Memory: {{}}

    Arguments:
    {{"customer_id": "BOFAMAST", "amount": 50000000}}

    Current Node: draft_email_2
    Tool Name: draft_email
    Tool Description: Drafts an email. Arguments: subject (required, string), recipient (required, string), content (required, string).
    Current Memory: {{"predict_payment_timing_2": "some_prediction_result"}}

    Arguments:
    {{
        "subject": "Payment Prediction",
        "recipient": "3333@ddd.com",
        "content": "The predicted business day for ABNMST is {{predict_payment_timing_2}}"
    }}

    Current Node: summarize_text_1
    Tool Name: summarize_text
    Tool Description: Summarizes the given text. Arguments: text (required, string)
    Current Memory: {{"predict_payment_timing_2": "some_prediction_result"}}

    Arguments:
    {{
        "text": "some_prediction_result"
    }}

    **Generate the arguments for the tool specified by the Current Node:**

    Arguments:
    """
        if feedback_prompt_json_error:
            prompt += f"""

        **Previous response was NOT valid JSON and caused a parsing error. Please generate a VALID JSON arguments response, taking into account the following feedback:**
        **Feedback on JSON Error:**
        {feedback_prompt_json_error}

        **Ensure the Revised Arguments are Valid JSON:**
        """
        arguments_str = self.llm_manager.llm_call(prompt, llm_type)

        try:
            arguments = json.loads(arguments_str)
        except json.JSONDecodeError as e:
            invalid_json_string = arguments_str # Capture invalid JSON string
            logging.error(f"JSON Parse Error (Attempt {retry_count + 1}) for node '{node}': {e}")
            logging.error(f"Invalid JSON String: {invalid_json_string}")

            if retry_count < max_retries:
                feedback_prompt_json_error = f"Your response was expected to be in valid JSON format, but it failed to parse with error: '{e}'. The invalid JSON string was: '{invalid_json_string}'. Please return ONLY valid JSON in your next response."
                logging.warning(f"Regenerating arguments for node '{node}' due to JSON parsing error. (Retry {retry_count + 1}/{max_retries})")
                return self._determine_tool_and_arguments_for_node( # Recursive retry
                    node, graph, memory_manager, tools, client, model_deployment_name, llm_type, retry_count + 1, max_retries, feedback_prompt_json_error # Increment retry count, pass feedback
                )
            else:
                logging.error(f"Maximum retries exceeded for argument regeneration after JSON parsing errors for node: '{node}'.")
                logging.error(f"Original JSON parsing error: {e}")
                logging.error(f"Last invalid JSON String: {invalid_json_string}")
                raise ValueError(f"Failed to generate valid JSON arguments for node '{node}' after multiple retries due to JSON parsing errors.") from e # Re-raise original exception or new ValueError

        return tool_name, arguments
    
    async def _regenerate_arguments_with_feedback(self, tool_name: str, tool_description: str, graph: Dict, memory_manager: MemoryManager, feedback: str, node: str, llm_type: str = None, error_feedback: str = None) -> Dict:
         """
         Uses the LLM to regenerate arguments for a specific tool based on user feedback.
         """
         # Use error_feedback if available, otherwise use the original feedback
         effective_feedback = error_feedback if error_feedback else feedback
         print(" effective feedback", effective_feedback, "for graph", graph, "Node----", graph[node])
         # Substitute placeholders in the prompt arguments
         substituted_results = memory_manager._substitute_placeholders(graph[node].get("input",{}))
         print(" substituted result" , substituted_results)
         if isinstance(graph[node],dict):
              graph[node]["input"] = substituted_results
         else:
              graph[node] = substituted_results
 

         prompt = f"""
         You are an AI assistant that revises arguments for a given tool based on user feedback.
 
         **Current Node:** {node}
         **Graph:** {json.dumps(graph)}
         **Tool Name:** {tool_name}
         **Tool Description:** {tool_description}
         **Current Memory (Results of Previous Steps):** {json.dumps(memory_manager.results)}
         **User Feedback:** {effective_feedback}
 
         **Task:**
         Generate revised arguments for the tool in JSON format, taking into account the user feedback.
 
         **Important Instructions:**
         - **The output must be a valid JSON string and nothing else. Do not include any introductory or explanatory text.**
         - Use the information from the current node name (e.g., entity names, parameters) to determine argument values.
         - Refer to the tool description to understand what arguments are required and their expected format (e.g., string, number, etc.).
         - Use the results of previous steps (available in the `Current Memory`) if the tool needs them as input. Reference previous results using the format `{{node_name}}` (e.g., `{{predict_payment_timing_BOFAMAST}}`).
         - **If no arguments are needed, return an empty JSON object {{}}.**
         - **Carefully consider the user's feedback when generating the revised arguments. The user feedback specifically indicates which part of the arguments need to be changed and what the new value should be.**
 
         **Example:**
 
         Current Node: draft_email_ddd@www.com
         Graph: {{"start": ["predict_payment_timing_BOFAMAST", "predict_payment_timing_ABNMST"], "predict_payment_timing_BOFAMAST": ["draft_email_ddd@www.com"], "predict_payment_timing_ABNMST": ["draft_email_333@222.com"]}}
         Tool Name: draft_email
         Tool Description: Drafts an email. Arguments: subject (required, string), recipient (required, string), content (required, string).
         Current Memory: {{"predict_payment_timing_BOFAMAST": "some_prediction_result"}}
         User Feedback: The recipient should be "test@example.com", not "ddd@www.com".
 
         Revised Arguments:
         {{
             "subject": "Payment Prediction",
             "recipient": "test@example.com",
             "content": "The predicted business day for BOFAMAST is {{predict_payment_timing_BOFAMAST}}"
         }}
 
         **Generate the revised arguments for the tool specified by the Current Node:**
 
         Revised Arguments:
         """
         print ("regenerate prompt--->", prompt)                                              
         arguments_str = self.llm_manager.llm_call(prompt, llm_type)
         print ("post regenerate prompt exec--->", arguments_str)   
         try:
             revised_arguments = json.loads(arguments_str)
             print("post json regen", revised_arguments)
         except json.JSONDecodeError as e:
             logging.error(f"Failed to parse JSON response from LLM while regenerating arguments for node '{node}' due to error {e}.")
             logging.info(f"The invalid JSON string was: {arguments_str}.")
             return {}
         
         logging.info(f"Successfully regenerated arguments for tool '{tool_name}' at node '{node}': {revised_arguments}")
         return revised_arguments

    async def _regenerate_arguments_with_AI_feedback(self, tool_name: str, tool_description: str, graph: Dict, memory_manager: MemoryManager, feedback: str, node: str, llm_type: str = None, error_feedback: str = None) -> Dict:
         """
         Uses the LLM to regenerate arguments for a specific tool based on user feedback.
         """
         # Use error_feedback if available, otherwise use the original feedback
         effective_feedback = error_feedback if error_feedback else feedback
         print(" effective feedback", effective_feedback, "for graph", graph, "Node----", graph[node])
         # Substitute placeholders in the prompt arguments
         substituted_results = memory_manager._substitute_placeholders(graph[node].get("input",{}))
         print(" substituted result" , substituted_results)
         if isinstance(graph[node],dict):
              graph[node]["input"] = substituted_results
         else:
              graph[node] = substituted_results

         previous_arguments = graph[node].get("input", {})  # Get previous arguments
 

         prompt = f"""
         You are an AI assistant that revises arguments for a given tool based on an AI feedback, focusing on correcting **data type errors**.
 
         **Current Node:** {node}
         **Graph:** {json.dumps(graph)}
         **Tool Name:** {tool_name}
         **Tool Description:** {tool_description}
         **Current Memory (Results of Previous Steps):** {json.dumps(memory_manager.results)}
         **Previous Arguments:** {json.dumps(previous_arguments)}
         **AI Feedback:** {effective_feedback}
         
 
         **Task:**
         Generate revised arguments for the tool in JSON format, taking into account the user feedback.

         **Specific Instructions based on Feedback:**
            - The feedback indicates: "{effective_feedback}".  **Pay close attention to the type of error mentioned in the feedback.**
            - **If the feedback indicates a data type error (e.g., "input is not a number", "invalid email format", "file format not correct"), ensure that the regenerated arguments have the correct data types as specified in the Tool Description.**
            - For the `multiply_numbers` tool, both `num1` and `num2` MUST be numbers.
 
         **Important Instructions:**
         - **The output must be a valid JSON string and nothing else. Do not include any introductory or explanatory text.**
         - Use the information from the current node name (e.g., entity names, parameters) to determine argument values.
         - Refer to the tool description to understand what arguments are required and their expected format (e.g., string, number, etc.).
         - Use the results of previous steps (available in the `Current Memory`) if the tool needs them as input. Reference previous results using the format `{{node_name}}` (e.g., `{{predict_payment_timing_BOFAMAST}}`).
         - **If no arguments are needed, return an empty JSON object {{}}.**
         - **Carefully consider the user's feedback when generating the revised arguments. The user feedback specifically indicates which part of the arguments need to be changed and what the new value should be.**
          **If, after considering the feedback, you determine that it's NOT possible to regenerate valid arguments for this tool that would resolve the issue (e.g., the problem is not with arguments but with the tool itself or the plan), then return a JSON with a special key `"action": "cannot_regenerate_arguments"` and an empty `arguments` object, like this: `{{ "action": "cannot_regenerate_arguments", "arguments": {{}} }}`.  Do not return empty `{{}}` directly in this case.**
        - **Otherwise, if you can regenerate valid arguments, return them in JSON format as before.**

 
         **Example:**
 
         Current Node: draft_email_ddd@www.com
         Graph: {{"start": ["predict_payment_timing_BOFAMAST", "predict_payment_timing_ABNMST"], "predict_payment_timing_BOFAMAST": ["draft_email_ddd@www.com"], "predict_payment_timing_ABNMST": ["draft_email_333@222.com"]}}
         Tool Name: draft_email
         Tool Description: Drafts an email. Arguments: subject (required, string), recipient (required, string), content (required, string).
         Current Memory: {{"predict_payment_timing_BOFAMAST": "some_prediction_result"}}
         User Feedback: The recipient should be "test@example.com", not "ddd@www.com".
 
         Revised Arguments:
         {{
             "subject": "Payment Prediction",
             "recipient": "test@example.com",
             "content": "The predicted business day for BOFAMAST is {{predict_payment_timing_BOFAMAST}}"
         }}
 
         **Generate the revised arguments for the tool specified by the Current Node:**
 
         Revised Arguments:
         """
         print ("regenerate prompt--->", prompt)                                              
         arguments_str = self.llm_manager.llm_call(prompt, llm_type)
         print ("post regenerate prompt exec--->", arguments_str)   
         try:
             revised_arguments = json.loads(arguments_str)
             print("post json regen", revised_arguments)
         except json.JSONDecodeError as e:
             logging.error(f"Failed to parse JSON response from LLM while regenerating arguments for node '{node}' due to error {e}.")
             logging.info(f"The invalid JSON string was: {arguments_str}.")
             return {}
         
         logging.info(f"Successfully regenerated arguments for tool '{tool_name}' at node '{node}': {revised_arguments}")
         return revised_arguments
# --- InstructionRetriever and DynamicExampleSelector ---
# (These are the missing classes you requested)

class InstructionRetriever:
    def __init__(self, embeddings_file: str, bm25_top_n: int = 5, alpha: float = 0.5):
        """
        Initializes the InstructionRetriever.

        Args:
            embeddings_file: Path to the pickle file containing embeddings, instructions, and BM25 index.
            api_key: Your Google AI Gemini API key.
            bm25_top_n: The number of top candidates to retrieve using BM25 before re-ranking.
            alpha: The weighting factor for combining BM25 and cosine similarity scores (0.0 = pure BM25, 1.0 = pure cosine).
        """
        try:
             if not os.path.exists(embeddings_file):
                print(f"Embeddings file not found at {embeddings_file}. Creating a new file.")
                with open(embeddings_file, "wb") as f:
                      pickle.dump({"instructions": [], "bm25": None, "index": None}, f)
            
             data = safe_load_pickle(embeddings_file)
             self.instructions = data["instructions"]  # List of instruction dictionaries
             self.bm25 = data["bm25"]
             self.index = data["index"]
             #genai.configure(api_key=api_key)
             self.model = genai.GenerativeModel('models/embedding-001')
             self.bm25_top_n = bm25_top_n
             self.alpha = alpha
             self.vectorizer = TfidfVectorizer()  # Initialize TF-IDF vectorizer for keyword-based search
            

                        # Fit the vectorizer on the instruction texts for BM25
             if self.instructions:
                instruction_texts = [json.dumps(instr, sort_keys=True) for instr in self.instructions]
                self.vectorizer.fit(instruction_texts)
                tokenized_instructions = [text.split(" ") for text in instruction_texts]
                self.bm25 = BM25Okapi(tokenized_instructions)         
        except FileNotFoundError as e:
            raise FileNotFoundError(f"Error: Embeddings file not found at {embeddings_file}.")

    def preprocess_text(self, text):
        # Convert text to lowercase
        text = text.lower()
        # Remove punctuation (keep spaces and alphanumeric characters)
        text = re.sub(r'[^\w\s]', '', text)
        return text
    
    def get_relevant_instructions(self, input_text: str, num_instructions: int, similarity_threshold: float, api_key: str) -> List[Dict]:

        genai.configure(api_key=api_key)

        if not self.instructions or self.index is None:
            print("Warning: Embeddings data not loaded.")
            return []
        
        # 1. Preprocess input text for keyword matching (if needed):
        preprocessed_input = self.preprocess_text(input_text)  # Assuming you have a preprocess_text method

        relevant_instructions = []
        for instruction in self.instructions:
            # 2. Check for keyword matches in the specified fields:
            fields_to_check = [
                instruction.get("Input", ""),
                instruction.get("Goal", ""),
                instruction.get("User Input/Goal", "")  # Assuming you add this field consistently
            ]

            # Check if any of the fields contain the input text (using basic string matching)
            if any(preprocessed_input in field.lower() for field in fields_to_check):
                relevant_instructions.append(instruction)
            # if an instruction is found based on keywords, no need to check other instructions
            if len(relevant_instructions) >= num_instructions:
                return relevant_instructions

        # 3. Use Embeddings if no keyword matches are found:
        if not relevant_instructions:
            print("No matching instructions found based on keywords. Using embeddings.")
            # Embedding Retrieval 
            input_embedding_response = genai.embed_content(
                model="models/embedding-001",
                content=[input_text],
                task_type="retrieval_query",
            )
            input_embedding = np.array(input_embedding_response['embedding'], dtype='float32')
            faiss.normalize_L2(input_embedding)
            D, I = self.index.search(input_embedding, k=len(self.instructions))
            embedding_similarities = D[0]

            embedding_results_with_scores = sorted(
                enumerate(embedding_similarities), key=lambda item: item[1], reverse=True
            )

            # Filter by Threshold and Retrieve Instructions
            for i, similarity in zip(I[0], embedding_similarities):
                if similarity >= similarity_threshold:
                    relevant_instructions.append(self.instructions[i])
                    if len(relevant_instructions) >= num_instructions:
                        break
            
        return relevant_instructions


                                                                                       
class DynamicExampleSelector:
    def __init__(self, embeddings_file: str, bm25_top_n: int = 10, alpha: float = 0.2):
        """
        Initializes the DynamicExampleSelector.

        Args:
            embeddings_file: Path to the pickle file containing embeddings, examples, and BM25 index.
            api_key: Your Google AI Gemini API key.
            bm25_top_n: The number of top candidates to retrieve using BM25 before re-ranking.
            alpha: The weighting factor for combining BM25 and cosine similarity scores (0.0 = pure BM25, 1.0 = pure cosine).
        """
        try:
            data = safe_load_pickle(embeddings_file)
            self.examples = data.get("examples", [])  # Load examples if they exist, otherwise set to an empty list
            self.bm25 = data["bm25"]
            self.index = data["index"]
            self.model = genai.GenerativeModel('models/embedding-001')
            self.bm25_top_n = bm25_top_n
            self.alpha = alpha  # Add alpha parameter for weighting
            self.vectorizer = TfidfVectorizer()  # Initialize TF-IDF vectorizer for keyword-based search

        except FileNotFoundError as e :
            raise FileNotFoundError(f"Error: Embeddings file not found at {embeddings_file}.")
       
    def preprocess_text(self, text):
        # Convert text to lowercase
        text = text.lower()
        # Remove punctuation
        text = re.sub(r'[^\w\s]', '', text)
        return text
    
    def get_relevant_examples(self, goal: str, example_type: str, good_or_bad: str, num_examples: int, similarity_threshold: float, api_key: str) -> List[str]:
        # Preprocess documents and goal
        preprocessed_goals = [self.preprocess_text(ex["goal"]) for ex in self.examples]
        preprocessed_goal = self.preprocess_text(goal)
        #print("Preprocessed Goals:", preprocessed_goals)
        #print("Preprocessed Goal:", preprocessed_goal)

        # Stage 1: Keyword-Based Search (BM25)
        X = self.vectorizer.fit_transform(preprocessed_goals)  # Use TF-IDF for sparse vectors
        tokenized_goal = preprocessed_goal.split(" ")
        bm25_scores = self.bm25.get_scores(tokenized_goal)
        #print("BM25 Scores:", bm25_scores)
        top_n_indices = np.argsort(bm25_scores)[::-1][:self.bm25_top_n]
        candidate_examples = [self.examples[i] for i in top_n_indices]
        #print("Top N Indices:", top_n_indices)
        #print("Candidate Examples:", candidate_examples)

        # Stage 2: Vector-Based Search (Cosine Similarity)
        genai.configure(api_key=api_key)
        goal_embedding_response = genai.embed_content(
            model="models/embedding-001",
            content=preprocessed_goal,            
            task_type="retrieval_query",
        )
        goal_embedding = np.array([goal_embedding_response['embedding']], dtype='float32')
        faiss.normalize_L2(goal_embedding)

        D, I = self.index.search(goal_embedding, k=len(candidate_examples))
        similarities = 1 - D[0]  # Convert distances to similarities
        #print("Distances (D):", D)
        #print("Indices (I):", I)
        #print("Similarities:", similarities)

        # Stage 3: Hybrid Scoring
        hybrid_scores = []
        for i, example in enumerate(candidate_examples):
            bm25_score = bm25_scores[top_n_indices[i]]
            vector_similarity = similarities[i]
            hybrid_score = (1 - self.alpha) * bm25_score + self.alpha * vector_similarity
            hybrid_scores.append((example, hybrid_score))
            #print(f"Example {i}: BM25 Score = {bm25_score}, Vector Similarity = {vector_similarity}, Hybrid Score = {hybrid_score}")

        # Re-rank candidates based on hybrid scores
        ranked_candidates = sorted(hybrid_scores, key=lambda x: x[1], reverse=True)
        #print("Ranked Candidates:", ranked_candidates)

        # Stage 4: Selection
        relevant_examples = []
        for example, hybrid_score in ranked_candidates:
            if example["example_type"] == example_type and example["good_or_bad_example"] == good_or_bad and hybrid_score >= similarity_threshold:
                relevant_examples.append(example["example"])
                if len(relevant_examples) >= num_examples:
                    break
        #print("Relevant Examples:", relevant_examples)

        return relevant_examples
    
    def insert_examples_into_prompt(self, prompt: str, goal: str, example_type: str, num_good_examples: int, num_bad_examples: int, similarity_threshold: float, api_key: str) -> str:
        """
        Inserts relevant examples into the prompt by replacing placeholders.

        Args:
            prompt: The prompt template string.
            goal: The user's goal.
            example_type: The type of example to insert.
            num_good_examples: The number of good examples to retrieve.
            num_bad_examples: The number of bad examples to retrieve.
            similarity_threshold: The minimum cosine similarity score for an example to be considered relevant.

        Returns:
            The modified prompt string with examples inserted.
        """
        good_examples = self.get_relevant_examples(goal, example_type, "good", num_good_examples, similarity_threshold, api_key)
        #bad_examples = self.get_relevant_examples(goal, example_type, "bad", num_bad_examples, similarity_threshold)
        #print("good examples", good_examples)
        good_examples_str = "\n".join(good_examples)
        #print("good examples", good_examples_str)
        #bad_examples_str = "\n".join(bad_examples)

        prompt = good_examples_str
        #prompt = prompt.replace(f"{{EXAMPLE_TYPE_{example_type.upper()}_BAD_EXAMPLES}}", bad_examples_str)
        #print("insert into", prompt)
        return prompt

# --- Instruction Manager ---
class InstructionManager:
    def __init__(self, embeddings_file: str, api_key: str):
        self.api_key = api_key  # Store API key here
        self.embeddings_file = embeddings_file
        self.DYNAMIC_EMBEDDINGS_FILE = DYNAMIC_EMBEDDINGS_FILE
        self.instruction_retriever = InstructionRetriever(embeddings_file)
        self.dynamic_example_selector = DynamicExampleSelector(DYNAMIC_EMBEDDINGS_FILE)

    def get_relevant_instructions(self, input_text: str, num_instructions: int, similarity_threshold: float) -> List[Dict]:
        genai.configure(api_key=self.api_key)
        return self.instruction_retriever.get_relevant_instructions(input_text, num_instructions, similarity_threshold, self.api_key)

    def insert_examples_into_prompt(self, prompt: str, goal: str, example_type: str, num_good_examples: int, num_bad_examples: int, similarity_threshold: float) -> str:
        genai.configure(api_key=self.api_key)
        return self.dynamic_example_selector.insert_examples_into_prompt(prompt, goal, example_type, num_good_examples, num_bad_examples, similarity_threshold,self.api_key)

    
    
    def create_embeddings(self,file_content: Union[str, BytesIO], embeddings_file: str):
        """
        Generates embeddings for instructions from a .docx file and saves them to a .pkl file.

        Args:
            instructions_file: Path to the .docx file containing instructions.
            embeddings_file: Path to the output .pkl file for storing embeddings, instructions, and indices.
            api_key: Your Google AI Gemini API key.

        Raises:
            HTTPException: If any error occurs during processing.
        """
        try:
            # 1. Extract Instructions from .docx
            try:
                if isinstance(file_content, str):
                    doc = Document(file_content)
                else:
                    doc = Document(file_content)
                #doc = Document(instructions_file)
                new_instructions = []
                current_instruction = {}
                reading_field = None  # Flag to track which field is being read (None = not reading multi-line)
                field_content = ""  # Accumulator for multi-line field content

                paragraph_iterator = iter(doc.paragraphs)  # Create an iterator
                while True:  # Use a while loop for paragraph iteration
                    try:
                        paragraph = next(paragraph_iterator)
                        text = paragraph.text.strip()
                        if not text:
                            continue

                        if text.startswith("Input:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            if current_instruction:
                                new_instructions.append(current_instruction)
                            current_instruction = {"Input": text.replace("Input:", "").strip()}
                        elif text.startswith("Goal:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            current_instruction["Goal"] = text.replace("Goal:", "").strip()
                        elif text.startswith("Tools to be used:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            tools_str = text.replace("Tools to be used:", "").strip()
                            current_instruction["Tools to be used"] = [tool.strip() for tool in tools_str.split(',')]
                        elif text.startswith("Sequence of Tools:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            sequence_str = text.replace("Sequence of Tools:", "").strip()
                            tool_sequence_list = [step.strip().split('.', 1)[-1].strip() for step in sequence_str.splitlines() if step.strip()] 
                            current_instruction["Sequence of Tools"] = tool_sequence_list # Store as list
                            #current_instruction["Sequence of Tools"] = sequence_str
                        elif text.startswith("Additional Information:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            current_instruction["Additional Information"] = text.replace("Additional Information:", "").strip()
                        elif text.startswith("User feedback for the plan:"): # User feedback
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            current_instruction["User feedback for the plan"] = text.replace("User feedback for the plan:", "").strip()
                        elif text.startswith("Revised Tools:"): # Revised Tools
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            current_instruction["Revised Tools"] = text.replace("Revised Tools:", "").strip()
                        elif text.startswith("Revised Sequence of Tools:"): # Revised Sequence
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = None
                            current_instruction["Revised Sequence of Tools"] = text.replace("Revised Sequence of Tools:", "").strip()
                        elif text.startswith("Graph:"):
                            self._store_field_content(current_instruction, reading_field, field_content) # Store accumulated content
                            reading_field = "Graph" # Start reading graph content, set reading_field to "Graph"
                            field_content = text.replace("Graph:", "").strip() # Initialize field_content
                        elif reading_field is not None: # Accumulate content if reading_field is set (for any field)
                            field_content += "\n" + text
                        elif text == "---" and not reading_field:  # Stop accumulating if "---" and not reading a field
                                reading_field = None
                                if current_instruction:
                                    new_instructions.append(current_instruction)
                                    current_instruction = {}
                        elif reading_field: # Accumulate content if reading_field is not None
                                field_content += "\n" + text
                        else:  # Handle text outside known keys
                                if current_instruction.get("Additional Information"):
                                    current_instruction["Additional Information"] += "\n" + text
                                elif current_instruction:  # if current instruction exists, append to the last field
                                    last_key = list(current_instruction.keys())[-1]
                                    current_instruction[last_key] += "\n" + text
                                else:
                                    pass  # or log warning about unexpected text

                    except StopIteration:  # Break loop when no more paragraphs
                        break

                # Process accumulated field content after loop
                if current_instruction and "Graph" == reading_field: # Process Graph if it was being read
                    graph_content_to_parse = field_content.strip() # Strip leading/trailing whitespace FIRST
                    # Aggressively remove trailing "---" if present (and any whitespace around it)
                    graph_content_to_parse = re.sub(r'\s*---\s*$', '', graph_content_to_parse, flags=re.MULTILINE)
                    logging.info(f"CREATE_EMBEDDINGS: Graph content before parsing for goal: {current_instruction.get('Goal')}:\n{field_content}") # <---- ADD THIS LOGGING
                    try:
                        current_instruction["Graph"] = json.loads(graph_content_to_parse)
                    except json.JSONDecodeError as e:
                        logging.error(f"CREATE_EMBEDDINGS: JSONDecodeError for goal: {current_instruction.get('Goal')}, Error: {e}")
                        current_instruction["Graph"] = graph_content_to_parse
                elif current_instruction and "Additional Information" == reading_field: # Process Additional Info if it was being read
                     current_instruction["Additional Information"] = field_content.strip()


                if current_instruction:
                    new_instructions.append(current_instruction)


            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Could not extract instructions: {e}")
            # 2. Configure Gemini API
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel('models/embedding-001')

            # 3. Load Existing Data (If File Exists)
            existing_instructions = []
            existing_bm25 = None
            existing_index = None
            try:
                if os.path.exists(embeddings_file):
                    try:
                        existing_data = safe_load_pickle(embeddings_file)
                        existing_instructions = existing_data["instructions"]
                        existing_bm25 = existing_data["bm25"]
                        existing_index = existing_data["index"]
                    except Exception as e:
                        raise HTTPException(status_code=500, detail=f"Could not load existing embeddings: {e}")
            except FileNotFoundError:
                raise FileNotFoundError(f"Embeddings file not found at {embeddings_file}. Creating a new file.")
            # 4. Merge New Instructions
            all_instructions = existing_instructions + new_instructions

            # 5. Create Embeddings
            instruction_texts = []
            for instruction in all_instructions:
                text = "\n".join([f"{field}: {value}" for field, value in instruction.items() if value])
                instruction_texts.append(text)

            # 6. Create BM25 Index
            tokenized_instructions = [text.split(" ") for text in instruction_texts]
            bm25 = BM25Okapi(tokenized_instructions)

            # 7. Generate Embeddings with Gemini API
            try:
                embeddings_response = genai.embed_content(
                    model="models/embedding-001",
                    content=instruction_texts,
                    task_type="retrieval_document",
                )
                embeddings = np.array(embeddings_response['embedding'], dtype='float32')
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error generating embeddings: {e}")

            # 8. Create FAISS Index
            index = faiss.IndexFlatIP(embeddings.shape[1])
            faiss.normalize_L2(embeddings)
            index.add(embeddings)

            # 9. Save Data to Pickle File
            try:
                with open(embeddings_file, "wb") as f:
                    pickle.dump({"instructions": all_instructions, "bm25": bm25, "index": index}, f)
                print(f"Embeddings, instructions, BM25 index, and FAISS index saved to {embeddings_file}")
                return "Embedding successful"
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error saving embeddings to file: {e}")

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {e}")
    
    def _store_field_content(self, current_instruction, reading_field, field_content):
        """Helper function to store accumulated field content."""
        if reading_field and current_instruction is not None:
            if reading_field == "Graph":
                try:
                    current_instruction[reading_field] = json.loads(field_content)
                except json.JSONDecodeError as e:
                    current_instruction[reading_field] = field_content # Store as string if JSON parsing fails
            else:
                current_instruction[reading_field] = field_content


    def create_embeddings_old(self,file_content: Union[str, BytesIO], embeddings_file: str):
        """
        Generates embeddings for instructions from a .docx file and saves them to a .pkl file.

        Args:
            instructions_file: Path to the .docx file containing instructions.
            embeddings_file: Path to the output .pkl file for storing embeddings, instructions, and indices.
            api_key: Your Google AI Gemini API key.

        Raises:
            HTTPException: If any error occurs during processing.
        """
        try:
            # 1. Extract Instructions from .docx
            try:
                if isinstance(file_content, str):
                    doc = Document(file_content)
                else:
                    doc = Document(file_content)
                #doc = Document(instructions_file)
                new_instructions = []
                current_instruction = {}

                for paragraph in doc.paragraphs:
                    text = paragraph.text.strip()
                    if not text:
                        continue

                    if text.startswith("Input:"):
                        if current_instruction:
                            new_instructions.append(current_instruction)
                        current_instruction = {"Input": text.replace("Input:", "").strip()}
                    elif text.startswith("Goal:"):
                        current_instruction["Goal"] = text.replace("Goal:", "").strip()
                    elif text.startswith("Tools to be used:"):
                        tools_str = text.replace("Tools to be used:", "").strip()
                        current_instruction["Tools"] = [tool.strip() for tool in tools_str.split(',')]
                    elif text.startswith("Sequence of Tools:"):
                        sequence_str = text.replace("Sequence of Tools:", "").strip()
                        current_instruction["Sequence"] = [step.strip() for step in sequence_str.split(',')]
                    elif text.startswith("Additional Information:"):
                        current_instruction["Additional Information"] = text.replace("Additional Information:", "").strip()
                    elif text.startswith("Graph:"):
                        try:
                            graph_str = text.replace("Graph:", "").strip()
                            current_instruction["Graph"] = json.loads(graph_str)
                        except json.JSONDecodeError:
                            current_instruction["Graph"] = graph_str # If it is not a valid JSON then store it as string
                    elif text == "---":
                        if current_instruction:
                            new_instructions.append(current_instruction)
                            current_instruction = {}

                if current_instruction:
                    new_instructions.append(current_instruction)

            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Could not extract instructions: {e}")

            # 2. Configure Gemini API
            genai.configure(api_key=self.api_key)
            model = genai.GenerativeModel('models/embedding-001')

            # 3. Load Existing Data (If File Exists)
            existing_instructions = []
            existing_bm25 = None
            existing_index = None        
            try:
                os.path.exists(embeddings_file)
                existing_data = safe_load_pickle(embeddings_file)
                existing_instructions = existing_data["instructions"]
                existing_bm25 = existing_data["bm25"]
                existing_index = existing_data["index"]
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Could not load existing embeddings: {e}")
            
            # 4. Merge New Instructions
            all_instructions = existing_instructions + new_instructions

            # 5. Create Embeddings
            instruction_texts = []
            for instruction in all_instructions:
                text = "\n".join([f"{field}: {value}" for field, value in instruction.items() if value])
                instruction_texts.append(text)

            # 6. Create BM25 Index
            tokenized_instructions = [text.split(" ") for text in instruction_texts]
            bm25 = BM25Okapi(tokenized_instructions)

            # 7. Generate Embeddings with Gemini API
            try:
                embeddings_response = genai.embed_content(
                    model="models/embedding-001",
                    content=instruction_texts,
                    task_type="retrieval_document",
                )
                embeddings = np.array(embeddings_response['embedding'], dtype='float32')
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error generating embeddings: {e}")

            # 8. Create FAISS Index
            index = faiss.IndexFlatIP(embeddings.shape[1])
            faiss.normalize_L2(embeddings)
            index.add(embeddings)

            # 9. Save Data to Pickle File
            try:
                with open(embeddings_file, "wb") as f:
                    pickle.dump({"instructions": all_instructions, "bm25": bm25, "index": index}, f)
                print(f"Embeddings, instructions, BM25 index, and FAISS index saved to {embeddings_file}")
                return "Embedding successful"
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"Error saving embeddings to file: {e}")

        except Exception as e:
            raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {e}")

 
    async def update_instructions_from_feedback(self):
        """Updates the instructions by reading the Feedback_Instructions.docx file and calling upload_instructions method."""
        try:
            # 1. Read the Feedback Instructions Document
            try:
                doc = Document("Feedback_Instructions.docx")
                # Write into memory, to avoid issues if file is being used by another process
                file_stream = BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                # Prepare a fake upload file
                class FakeUploadFile:
                    def __init__(self, file_stream, filename):
                        self.file = file_stream
                        self.filename = filename

                    async def read(self):
                        return self.file.read()

                file = FakeUploadFile(file_stream, filename="Feedback_Instructions.docx")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading Feedback Instructions document: {e}")
        
            # 2. Trigger upload_instructions
            # result = await upload_instructions(file=file)
            
            # return {"message": f"Instructions updated successfully: {result}"}

        except HTTPException as e:
            raise  # Re-raise HTTP exceptions from create_embeddings
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"An unexpected error occurred: {e}")

# --- Feedback Manager ---
class FeedbackManager:
    def __init__(self, feedback_log_directory: str, llm_manager: LLMManager,send_message_func):
        self.feedback_log_directory = feedback_log_directory
        self.llm_manager = llm_manager  # Store the llm_manager instance
        self.send_message = send_message_func
        os.makedirs(self.feedback_log_directory, exist_ok=True)

    def _log_feedback(self, plan_id: str, goal: str, old_plan: List, feedback: str, iteration: int, success: bool):
        """Logs feedback about a plan."""
        timestamp = datetime.now().isoformat()
        feedback_entry = {
            "timestamp": timestamp,
            "plan_id": plan_id,
            "goal": goal,
            "old_plan": old_plan,
            "feedback": feedback,
            "iteration": iteration,
            "success": success
        }
        file_path = os.path.join(self.feedback_log_directory, "feedback_log.json")
        try:
            if os.path.exists(file_path):
                with open(file_path, "r") as f:
                    try:
                        existing_data = json.load(f)
                    except json.JSONDecodeError:
                        existing_data = []
                existing_data.append(feedback_entry)
                with open(file_path, "w") as f:
                    json.dump(existing_data, f, indent=4)
            else:
                with open(file_path, "w") as f:
                    json.dump([feedback_entry], f, indent=4)
            logging.info(f"Feedback logged to: {file_path}")
        except Exception as e:
            logging.error(f"Error logging feedback: {e}")


    def _write_feedback_to_word(self, request, original_plan, revised_plan=None, feedback=None, llm_type: str = None):
        """Writes feedback to a Word document with structured format."""
        doc_path = "Feedback_Instructions.docx"
        instant_doc_path = "Feedback_Instructions_instant.docx"
        print("Inside write Feedback--->", request, original_plan, revised_plan, feedback)

        # Open or create the main feedback document
        try:
            doc = Document(doc_path)
        except Exception:
            doc = Document()

        # Create or overwrite the instant feedback document
        instant_doc = Document()

        # Add Separator
        doc.add_paragraph("---")
        instant_doc.add_paragraph("---")

        if request.unstruct_input:
            doc.add_paragraph(f"Input: {request.unstruct_input}")
            instant_doc.add_paragraph(f"Input: {request.unstruct_input}")
        doc.add_paragraph(f"Goal: {request.goal}")
        instant_doc.add_paragraph(f"Goal: {request.goal}")

        # ---Original Plan Tools ---
        original_tools_names = [node.get("tool") for node in original_plan.get("graph", {}).values() if isinstance(node, dict) and node.get("tool")]
        original_tools_str = ", ".join(original_tools_names)
        doc.add_paragraph(f"Tools to be used: {original_tools_str}")
        instant_doc.add_paragraph(f"Tools to be used: {original_tools_str}")

        # ---Original Plan Sequence of Tools ---
        sequence_nodes = self._get_execution_sequence_from_graph(original_plan["graph"])
        sequence_str = ""
        for idx, node in enumerate(sequence_nodes):
            sequence_str += f"{idx+1}. {node}"
        doc.add_paragraph(f"Sequence of Tools:{sequence_str}")
        instant_doc.add_paragraph(f"Sequence of Tools:{sequence_str}")

        # ---Original Plan Graph ---
        original_graph = original_plan.get("graph", {}) # Get graph from original plan, default to empty dict if missing
        graph_json_str = json.dumps(original_graph, indent=4) # Format graph to JSON string
        doc.add_paragraph("Graph:").add_run(f"\n{graph_json_str}").font.name = "Courier New" # Add Graph with code style font
        instant_doc.add_paragraph("Graph:").add_run(f"\n{graph_json_str}").font.name = "Courier New" # Add Graph with code style font

        if revised_plan:
            # ---Revised Plan Tools ---
            revised_tools_names = [node.get("tool") for node in revised_plan.get("graph", {}).values() if isinstance(node, dict) and node.get("tool")]
            revised_tools_str = ", ".join(revised_tools_names)

            # ---Revised Plan Sequence of Tools ---
            revised_sequence_nodes = self._get_execution_sequence_from_graph(revised_plan["graph"])
            revised_sequence_str = ""
            for idx, node in enumerate(revised_sequence_nodes):
                revised_sequence_str += f"{idx+1}. {node}\n"

            instruction_content = self.llm_manager._generate_instruction_content_with_llm(request, original_tools_names, sequence_nodes, revised_tools_names, revised_sequence_nodes, feedback, llm_type=llm_type)
        else:
            instruction_content = self.llm_manager._generate_instruction_content_with_llm(request, original_tools_names, sequence_nodes, feedback=feedback, llm_type=llm_type)

        if feedback:
            doc.add_paragraph(f"User feedback for the plan: {feedback}")
            instant_doc.add_paragraph(f"User feedback for the plan: {feedback}")

        if revised_plan:
            # ---Revised Plan Tools ---
            doc.add_paragraph(f"Revised Tools: {revised_tools_str}")
            instant_doc.add_paragraph(f"Revised Tools: {revised_tools_str}")

            # ---Revised Plan Sequence of Tools ---
            sequence_str = ""
            for idx, node in enumerate(revised_sequence_nodes):
                sequence_str += f"{idx+1}. {node}\n"
            doc.add_paragraph(f"Revised Sequence of Tools:\n{revised_sequence_str}")
            instant_doc.add_paragraph(f"Revised Sequence of Tools:\n{revised_sequence_str}")

        doc.add_paragraph(f"Additional Information: {instruction_content}")
        instant_doc.add_paragraph(f"Additional Information: {instruction_content}")

        doc.add_paragraph("---")
        instant_doc.add_paragraph("---")

        # Save the documents
        doc.save(doc_path)
        instant_doc.save(instant_doc_path)

        print("Feedback instructions saved")

                                                                                                                                                                                                                   
    def _get_execution_sequence_from_graph(self, graph: Dict) -> List:
       """
         Extracts the sequence of execution from the graph by following the next attributes.

         Args:
             graph: A dictionary with graph structure.

         Returns:
              A list of nodes in the order they should be executed.
        """
       executed_nodes = set() # Add a set to store executed nodes
       current_nodes = graph.get("start",[])
       output = []
       while current_nodes:
            next_nodes = []
            for node in current_nodes:
               if node == "start" or node == "end" or node in executed_nodes:
                    continue
               output.append(node)
               if graph[node].get("next"):
                  for next_node in graph[node].get("next", []):
                      if next_node not in executed_nodes and next_node != "end":
                         next_nodes.append(next_node)
               executed_nodes.add(node) # Add node to executed set
            current_nodes = next_nodes # set the next node
       return output

# --- Agent ---
class Agent:
    """
    Represents an AI agent that can generate plans and execute them using a set of tools.
    Now supports both Azure OpenAI and Google Gemini models.
    """

    def __init__(self, config):
        """
        Initializes the Agent with necessary configurations.
        """
        self.config = config
        self.llm_manager = LLMManager(agent=self,
            api_key=config["api_key"],
            endpoint=config["endpoint"],
            model_deployment_name=config["model_deployment_name"],
            gemini_api_key=config["gemini_api_key"],
            gemini_model_name=config["gemini_model_name"],
            default_llm=config.get("default_llm", "gemini")
        )
        self.tool_executor = ToolExecutor(config["tool_config_file"], self.llm_manager)
        self.memory_manager = MemoryManager()
        self.plan_manager = PlanManager(self.llm_manager, self,send_message_func=self.send_message)
        self.argument_manager = ArgumentManager(self.llm_manager,self.memory_manager,send_message_func=self.send_message)
        self.instruction_manager = InstructionManager(config.get("instruction_embeddings_file", "instruction_embeddings.pkl"), config["gemini_api_key"])
        self.feedback_manager = FeedbackManager(config["feedback_log_directory"], self.llm_manager,send_message_func=self.send_message) # Pass llm_manager
        self.plan_storage: Dict[str, Dict] = {}
        self.pending_prompts = {}  # Dictionary to track pending prompts


    async def run_agent(self, request: RunAgentRequest, session_id: str):
        """Generates a plan for the given goal and executes it."""
        
        #Check if the goal is to clear memory
        if request.goal and request.goal.lower() == "clear memory":
            self.memory_manager.clear_memory()
            self.memory_manager.clear_results()
            self.memory_manager.clear_conversation_history()
            self.plan_storage.clear()
            # #await self.send_message(message_type=MessageType.MESSAGE.value, message="Memory cleared.", session_id=session_id)
            return
       

        print("Request received--->", request)

        print(f"AGENT.RUN_AGENT: Session ID received in run_agent: {session_id}")

        # Reset the agent's memory at the beginning of each run
        self.memory_manager.clear_memory()
        self.memory_manager.clear_results()
        self.plan_storage.clear()

        # Store user goal in results as short-term memory
        self.memory_manager.results["user_goal"] = request.goal if request.goal else request.unstruct_input # Or just request.goal, depending on what you want to store
        logging.info(f"User goal stored in results: {request.goal or request.unstruct_input}")
         

        llm_type = request.llm_type or self.llm_manager.default_llm

        # 1. Get the previous conversation context
        previous_context = self.memory_manager.get_conversation_history()
        previous_context_str = json.dumps(previous_context, indent=2)

        # 2. Determine the user's input (goal or unstructured text)
        user_input = request.goal or request.unstruct_input
        self.memory_manager.add_to_conversation({"user_input": user_input})

        feedback = None # Declare the feedback variable here
        plan_feedback = None
        revised_plan = None
            # --- Goal Inference Logic ---
        #goals = []  # Initialize goals to an empty list
        inferred_goals = []  # Initialize an empty list for inferred goals
        if request.goal:
            # Use the provided goal directly
            goals = [request.goal]
        elif request.unstruct_input:
            print("Inferring goal from unstruct_input ",request.unstruct_input)
            log_message = f"Inferring the goal from the Input..."
            # #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)
            # Infer goal from email content or ticket description
            if request.split_execution:
                inferred_goals = self.llm_manager.infer_goal_from_text_Split(request.unstruct_input, llm_type,self.tool_executor.tools, instruction_retriever=self.instruction_manager.instruction_retriever,split_execution=request.split_execution, selected_agents = selected_agents,previous_context=previous_context_str)
                log_message = f"Inferred goal {inferred_goals}","session_id"
                # #await self.send_message#await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id) # Pass log_message dictionary as message
            else:
                inferred_goals = self.llm_manager.infer_goal_from_text(request.unstruct_input, llm_type,self.tool_executor.tools, instruction_retriever=self.instruction_manager.instruction_retriever,split_execution=request.split_execution, selected_agents=selected_agents,previous_context=previous_context_str)
                log_message = f"Inferred goal {inferred_goals}", "session_id"
                #await self.send_message(message_type=MessageType.LOG.value, message=log_message, session_id=session_id)            
                #inferred_goals = infer_goal_from_text(request.unstruct_input, llm_type,agent.tools, instruction_retriever=agent.instruction_retriever,split_execution=request.split_execution)
            #inferred_goals = infer_goal_from_text(request.unstruct_input, llm_type,agent.tools, instruction_retriever=agent.instruction_retriever,split_execution=request.split_execution)
            print(f"Inferred goal: {inferred_goals}")
            if  inferred_goals:
                print("entering here")
                goals = inferred_goals
                self.memory_manager.add_to_conversation({"inferred_goal": inferred_goals})
            else:            #Handle goal inference failure (e.g., return an error message)
                raise HTTPException(status_code=400, detail="Could not infer goal from input.")
        else:
            raise HTTPException(status_code=400, detail="A goal or user input must be provided.")
        # --- End of Goal Inference Logic ---

        # Get selected agents from the request
        selected_agents = request.agents if request.agents else ["All"]
        print("selected_agents", selected_agents)

        # --- Iterate through goals ---
        all_results = []  # List to accumulate results for all goals
        for goal in goals:
            print("goal: ", goal, "goals:", goals)
            #self.memory_manager.store_in_memory("user_goal", goal)
            selected_tools = None  # Initialize selected_tools
            plan = None
            feedback_tool = None
            plan_modified = False  # Flag to track if the plan was modified
            revised_plan = None # Initialize revised_plan

            input_goal = f" Input goal entered : {goal}"
            print(f"message_type={MessageType.LOG.value}, message={input_goal}, session_id={session_id}")
            #await self.send_message(message_type=MessageType.LOG.value, message=input_goal, session_id=session_id)

            if request.tool_verifier:
                selected_tools = self.llm_manager._select_tools(goal, self.tool_executor.tools, llm_type,selected_agents)
                # formatted_tools_list = self._format_toollist_for_ui(selected_tools)
                print("Selected Tools:")
                for tool in selected_tools:
                    print(f"- {tool['name']}: {tool['description']}")
                formatted_tools_list=[]
                feedback_tool = await self.get_user_input(
                    f" <b style='color: teal;'> List of skills chosen by AI for Execution:</b style='color: teal;'> <br><br> {formatted_tools_list} <br><br> Provide <b><span style='color: teal;'>feedback</span></b> to revise the Skill choices, sequence or type <b style='color: teal;'>'yes'</b style='color: teal;'> to proceed: ","tool_selection", session_id=session_id)

                while feedback_tool and feedback_tool.lower() != 'yes':
                    revised_tools = self.revise_tools(selected_tools, goal, feedback_tool, llm_type, selected_agents)
                    # formatted_rev_tools_list = self._format_toollist_for_ui(revised_tools)
                    print("Revised Tools based on feedback:")
                    for tool in revised_tools:
                        print(f"- {tool['name']}: {tool['description']}")
                    selected_tools = revised_tools
                    formatted_rev_tools_list=[]
                    feedback_tool = await self.get_user_input(
                        f" <b style='color: teal;'> Revised List of skills chosen by AI for Execution:</b style='color: teal;'> <br><br> {formatted_rev_tools_list} <br><br> Provide <b><span style='color: teal;'>feedback</span></b> to revise the Skill choices, sequence or type <b style='color: teal;'>'yes'</b style='color: teal;'> to proceed: ","tool_selection", session_id=session_id)

                # Generate the plan after tool selection and feedback
                plan = await self.plan_manager.generate_plan(goal, self.memory_manager,selected_tools, llm_type=llm_type, auto_verifier=request.auto_verifier,selected_agents=selected_agents,session_id=session_id)
                plan_modified = True  # Plan is modified after tool selection
                plan_id = plan.get("plan_id")

            if not plan:  # Generate plan only if not already generated in tool_verifier
                print("entering 2")
                plan = await self.plan_manager.generate_plan(goal, self.memory_manager, llm_type=llm_type, auto_verifier=request.auto_verifier,selected_agents=selected_agents,session_id=session_id)
                plan_id = plan.get("plan_id")
            
            original_plan = plan.copy() # Capture original plan before verification

            # --- Plan Verification ---
            if request.plan_verifier:
                if plan_modified:
                    # If plan was modified by tool verifier, re-prompt for verification
                    plan, goal, plan_id, plan_feedback = await self.verify_and_regenerate_plan(plan, goal, llm_type, lambda p, **kwargs: self.get_user_input(p,"plan_verification", **kwargs), session_id=session_id,selected_agents=selected_agents)
                else:
                    # Otherwise, proceed with normal verification
                    plan, goal, plan_id, plan_feedback = await self.verify_and_regenerate_plan(plan, goal, llm_type, lambda p, **kwargs: self.get_user_input(p, **kwargs), session_id=session_id,selected_agents=selected_agents)

                plan_id = plan.get("plan_id")
                # Update plan storage with the latest plan, goal, and plan_id after verification
                self.plan_manager.plan_storage[plan_id] = {"graph": plan.get("graph"), "goal": goal, "iteration": 0}
                if plan != original_plan:
                    revised_plan = plan.copy()  # Capture revised plan after plan verification
            else:            

                plan_id = plan.get("plan_id")

                self.plan_manager.plan_storage[plan_id] = {"graph": plan.get("graph"), "goal": goal, "iteration": 0}

                    
            


            if request.plan_verifier == True: # if plan verifier is set, then use plan_feedback
                self.feedback_manager._write_feedback_to_word(request, original_plan, revised_plan=revised_plan, feedback=plan_feedback,llm_type=llm_type)
            elif request.tool_verifier == True:  # if tool verifier is set, then use the feedback
                self.feedback_manager._write_feedback_to_word(request, original_plan, revised_plan=revised_plan, feedback=feedback_tool,llm_type=llm_type)  #Pass request, original_plan, revised_plan and feedback
            else:
                self.feedback_manager._write_feedback_to_word(request, original_plan, revised_plan=revised_plan, feedback=None,llm_type=llm_type)
    
            self.feedback_manager._log_feedback(plan_id, goal, plan, "", 0, success=True)
    
            print("auto learning", request.auto_learning)

        return {"results": plan}  # Return all results
               

    async def send_message(self, message_type: str, message: Any, request_id: str = None, prompt_type: str = None, session_id: str = None):
        """Sends a message to the websocket."""
        #message_json = {"type": message_type, "message": message}
        #message_json = message 
        message_json = {"type": message_type, "message": message, "session_id": session_id}
        if request_id:
            message_json["request_id"] = request_id  # Add request_id to the message
            if prompt_type:
                message_json["prompt_type"] = prompt_type  # Add prompt_type to the message

        print(f"AGENT.SEND_MESSAGE: Session ID being used: {session_id}")
        


    async def get_user_input(self, prompt: str, prompt_type: str = "general",session_id: str = None) -> str:
        """Sends a prompt to the UI and waits for user input."""
        request_id = str(uuid.uuid4())  # Generate a unique ID for this request
        logging.info(f"Sending prompt with request_id: {request_id}, prompt_type: {prompt_type}")  # Log the request_id and prompt_type
        print(f"Sending prompt with request_id: {request_id}, prompt_type: {prompt_type}, session_id: {session_id} ")
        self.pending_prompts[request_id] = asyncio.Future()  # Create a Future to track the response
        #prompt_message_json = {"type": MessageType.PROMPT.value, "message": prompt, "session_id": session_id, "request_id": request_id, "prompt_type": prompt_type} 
        ##await self.send_message_prompt(message_type=MessageType.PROMPT.value, message=prompt_message_json, session_id=session_id, request_id=request_id, prompt_type=prompt_type) # Keyword args
        #await self.send_message(message_type=MessageType.PROMPT.value, message=prompt, session_id=session_id, request_id=request_id, prompt_type=prompt_type)

        try:
            user_input = await asyncio.wait_for(self.pending_prompts[request_id], timeout=600)  # Wait for the Future to be resolved, with a timeout
            del self.pending_prompts[request_id]  # Clean up
            
            # --- ADD THIS: Store the user input in the conversation history ---
            self.memory_manager.add_to_conversation({
                "user_input": user_input,
                "prompt_type": prompt_type,  #  Good to store the type of prompt
                "request_id": request_id   #  And the request ID for debugging
            })
            # --------------------------------------------------------------------

            return user_input
        except asyncio.TimeoutError:
            logging.error(f"Timeout waiting for user input for request: {request_id}")
            del self.pending_prompts[request_id]  # Clean up on timeout
            #timeout_error_message_json = {"type": MessageType.ERROR.value, "message": "Timeout waiting for user input. Please try again.", "session_id": session_id, "request_id": request_id} # Create error_message_json dictionary
            ##await self.send_message(message_type=MessageType.ERROR.value, message=timeout_error_message_json, session_id=session_id, request_id=request_id) # Keyword args
            #await self.send_message(message_type=MessageType.ERROR.value, message="Timeout waiting for user input. Please try again.", session_id=session_id, request_id=request_id) # Pass ERROR STRING as 'message'
            return ""  # Re-raise the timeout error or handle it as needed


    
    def format_user_input_prompt(self,arguments, tool_name=None):
        formatted_args = json.dumps(arguments, indent=2).replace('\n', '<br>').replace(' ', '&nbsp;')
        return (
            # Increase max-width and add word-wrap properties
            f"<div style='background-color: #f8f9fa; padding: 15px; border-radius: 8px; "
            f"border: 1px solid #e9ecef; font-family: Arial, sans-serif; max-width: 1200px; " # Changed from 800px to 1200px
            f"word-wrap: break-word; overflow-wrap: break-word;'>" # Added these properties
            
            # Add similar word-wrap properties to nested divs
            f"<div style='margin-bottom: 12px; word-wrap: break-word; overflow-wrap: break-word;'>"
            f"<span style='background-color: #ff4757; color: white; padding: 3px 8px; "
            f"border-radius: 4px; font-weight: bold;'>Executing Skill</span>&nbsp;"
            f"<span style='background-color: #2196F3; color: white; padding: 3px 8px; "
            f"border-radius: 4px; font-weight: bold;'>{tool_name}</span>"
            f"</div>"
            
            # Add word-wrap to arguments div
            f"<div style='margin-bottom: 12px; word-wrap: break-word; overflow-wrap: break-word;'>"
            f"<div style='font-weight: bold; color: #2d3436; margin-bottom: 6px;'>Arguments:</div>"
            f"<div style='background-color: #f1f8ff; padding: 10px; border-radius: 4px; "
            f"font-family: monospace; color: #0366d6; white-space: pre-wrap;'>{formatted_args}</div>" # Added white-space: pre-wrap
            f"</div>"
            
            # Rest of the code remains the same
            f"<div style='margin-top: 15px; color: #2d3436;'>"
            f"Confirm or provide feedback:"
            f"<ul style='margin-top: 8px; margin-bottom: 0;'>"
            f"<li>Type <span style='background-color: #00b894; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>yes</span> to confirm</li>"
            f"<li>Provide <span style='background-color: #00b894; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>feedback</span> to revise values</li>"
            f"<li>Type <span style='background-color: #636e72; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>quit</span> to exit</li>"
            f"</ul>"
            f"</div>"
            f"</div>"
        )

    """def format_user_input_prompt(self,tool_name, arguments):
        formatted_args = json.dumps(arguments, indent=2).replace('\n', '<br>').replace(' ', '&nbsp;')
        return (
            f"<div style='background-color: #f8f9fa; padding: 15px; border-radius: 8px; "
            f"border: 1px solid #e9ecef; font-family: Arial, sans-serif; max-width: 800px;'>"
            f"<div style='margin-bottom: 12px;'>"
            f"<span style='background-color: #ff4757; color: white; padding: 3px 8px; "
            f"border-radius: 4px; font-weight: bold;'>Executing Skill</span>&nbsp;"
            f"<span style='background-color: #2196F3; color: white; padding: 3px 8px; "
            f"border-radius: 4px; font-weight: bold;'>{tool_name}</span>"
            f"</div>"
            f"<div style='margin-bottom: 12px;'>"
            f"<div style='font-weight: bold; color: #2d3436; margin-bottom: 6px;'>Arguments:</div>"
            f"<div style='background-color: #f1f8ff; padding: 10px; border-radius: 4px; "
            f"font-family: monospace; color: #0366d6;'>{formatted_args}</div>"
            f"</div>"
            f"<div style='margin-top: 15px; color: #2d3436;'>"
            f"Confirm or provide feedback:"
            f"<ul style='margin-top: 8px; margin-bottom: 0;'>"
            f"<li>Type <span style='background-color: #00b894; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>yes</span> to confirm</li>"
            f"<li>Provide <span style='background-color: #00b894; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>feedback</span> to revise values</li>"
            f"<li>Type <span style='background-color: #636e72; color: white; padding: 2px 6px; "
            f"border-radius: 3px; font-weight: bold;'>quit</span> to exit</li>"
            f"</ul>"
            f"</div>"
            f"</div>"
        )"""


    async def verify_arguments(self, tool_name: str, arguments: Dict, node: str, graph: Dict, memory_manager: MemoryManager, llm_type: str = None, get_user_input_func=None, session_id: str = None) -> bool:
        """
        Presents the tool name and generated arguments to the user for verification and allows for feedback-driven revisions.

        Args:
            tool_name: The name of the tool.
            arguments: The arguments generated for the tool.
            node: The current node in the graph.
            graph: The execution graph.
            memory: The current memory state.
            results: The results of previous steps.

        Returns:
            True if the user confirms the arguments are correct, False otherwise.
        """
        logging.info(f"Verifying arguments for tool: {tool_name} at node: {node} with arguments: {arguments} ")
        print(json.dumps({"type": "log", "message": f"Verifying arguments for tool: {tool_name} at node: {node} with arguments: {json.dumps(arguments, indent=2)}"}))
        print(f"\nTool: {tool_name}")
        print(f"Arguments: {json.dumps(arguments, indent=2)}")  # Pretty-print the arguments

        while True:
            #feedback = await get_user_input_func(f"<b><span style='color: red;'>Executing Skill...</span></b><b><span style='color: blue;'>{tool_name}</span></b> <br> Arguments:<span style='color: blue;'> {json.dumps(arguments, indent=2)}</span> <br> Are these values correct? If yes, type <b><span style='color: teal;'>'yes'</span></b> If no, provide <b><span style='color: teal;'>feedback</span></b> to revise the values or type 'quit' to exit: ")
            feedback = await get_user_input_func(self.format_user_input_prompt(arguments,tool_name),session_id=session_id)

            if feedback.lower() == "yes":
                logging.info(f"Arguments for tool '{tool_name}' at node '{node}' verified successfully.")
                return True

            elif feedback.lower() in ["quit", "q"]:
                logging.info("Exiting the program as per user request during argument verification.")
                print("Exiting the program.")
                sys.exit()

            else:  # User provided feedback
                logging.info(f"Feedback received: {feedback}. Revising the arguments...")
                # Store argument feedback in results (using node-specific key)
                self.memory_manager.results[f"argument_feedback_{node}"] = feedback # e.g., "argument_feedback_draft_email_1"
                logging.info(f"Argument feedback for node '{node}' stored in results: {feedback}")
                print("Feedback noted. Revising the arguments...")

                revised_arguments = await self.argument_manager._regenerate_arguments_with_feedback(
                    tool_name,
                    self.tool_executor.tools[tool_name]["description"],
                    graph,
                    memory_manager,
                    feedback,
                    node,
                    llm_type
                )

                if revised_arguments:
                    logging.info(f"Revised arguments for tool '{tool_name}': {revised_arguments}")
                    print(f"Revised Arguments: {json.dumps(revised_arguments, indent=2)}")

                    """confirmation = await get_user_input_func(
                        f"<b><span style='color: grey;'>Revised Values...</span></b><b><span style='color: blue;'>{revised_arguments}</span></b> <br> Are these values correct? If yes, type <b><span style='color: teal;'>'yes'</span></b> If no, provide <b><span style='color: teal;'>feedback</span></b> to revise the values or type 'quit' to exit: "
                    )"""

                    confirmation = await get_user_input_func(self.format_user_input_prompt(revised_arguments, tool_name),session_id=session_id)


                    if confirmation.lower() == "yes":
                        # Update the arguments for the current node in the graph
                        for node_item in graph:
                            if node_item == node:
                                if isinstance(graph[node], dict):
                                    if "input" in graph[node]:
                                        graph[node]["input"].clear()
                                        graph[node]["input"].update(revised_arguments)
                                        arguments.update(revised_arguments)  # Update the local 'arguments' as well
                                    else:
                                        graph[node]["input"] = revised_arguments
                                        arguments.update(revised_arguments) # Update the local 'arguments' as well
                                else:
                                    arguments.clear()
                                    arguments.update(revised_arguments)
                                break

                        logging.info(f"Arguments for tool '{tool_name}' at node '{node}' updated in the graph.")
                        print(f"Arguments for tool '{tool_name}' at node '{node}' updated in the graph.")
                        return True  # Arguments are correct after revision

                    elif confirmation.lower() in ["quit", "q"]:
                        logging.info("Exiting the program as per user request during argument verification.")
                        print("Exiting the program.")
                        sys.exit()

                    else:  # User provided feedback again
                        logging.info(f"Feedback received: {confirmation}. Revising the arguments again...")
                        print("Feedback noted. Revising the arguments again...")
                        arguments = revised_arguments  # Update arguments for the next iteration
                        continue # go to top of loop and re-prompt for feedback on revised arguments

                else:
                    logging.error(f"Failed to revise arguments for tool '{tool_name}' at node '{node}'.")
                    print("Failed to revise the arguments. Please try again or type 'quit' to exit.")

    
    
    async def verify_and_regenerate_plan(self, plan: Dict, goal: str, llm_type: str, get_user_input_func=None, session_id=None,selected_agents: List[str] = None) -> Dict:
        """
        Verifies the generated graph with the user. If the user provides feedback,
        regenerates the goal and the plan based on the feedback.
        """
        plan_id = plan.get("plan_id")
        plan_feedback = None
        while True:
            print("\nGenerated Graph:<br>")
            formatted_graph = ''
            # formatted_graph = self._format_graph_for_ui(plan.get("graph"))
            logging.info(f"Generated Graph:{json.dumps(plan.get('graph'), indent=2)}")
            feedback_plan = await get_user_input_func(
                f"<div style='margin-bottom: 16px; padding-bottom: 12px; border-bottom: 1px solid #e2e8f0;'>"
                f"<div style='color: #718096; font-size: 13px; margin-bottom: 4px; text-transform: uppercase; letter-spacing: 0.05em;'>AI Planning System</div>"
                #f"<span style='color: #16537E; font-size: 16px; font-weight: 700;'>Execution Graph - Step-by-Step Plan</span>"
                f"<span style='background-color: #3D85C6; color: white; padding: 3px 8px; "
                f"border-radius: 4px; font-weight: bold;'>Execution Plan - Step-by-Step Plan</span>&nbsp;"
                f"</div>"
                f"{formatted_graph}"
                # Enhanced feedback prompt section with blue theme
                f"<div style='margin: 16px -15px -12px -15px; padding: 12px 15px; background-color: #EBF8FF; border-top: 1px solid #e2e8f0; border-bottom-left-radius: 4px; border-bottom-right-radius: 4px;'>"
                f"<div style='display: flex; align-items: center; gap: 8px;'>"
                f"<span style='color: #0284C7; font-size: 20px;'>💡</span>"  # Light bulb icon
                f"<span style='color: #0C4A6E; font-weight: 600;'>Your Input Needed:</span>"
                f"</div>"
                f"<div style='margin-top: 8px; color: #0369A1;'>"
                f"Is this graph correct? If yes, type <span style='background-color: #DBEAFE; padding: 2px 8px; border-radius: 3px; font-weight: 600; color: #2563EB;'>yes</span>"
                f". If no, provide <span style='background-color: #DBEAFE; padding: 2px 8px; border-radius: 3px; font-weight: 600; color: #2563EB;'>feedback</span>"
                f" to revise the graph or type <span style='background-color: #DBEAFE; padding: 2px 8px; border-radius: 3px; font-weight: 600; color: #2563EB;'>quit</span> to exit."
                f"</div>"
                f"</div>"
                f"</div>",session_id=session_id)
            if feedback_plan.lower() == "yes":
                if plan_id:
                    self.plan_manager.plan_storage[plan_id] = {"graph": plan.get("graph"), "goal": goal, "iteration": 0}
                return plan, goal, plan_id, plan_feedback  # Return after plan is verified
            elif feedback_plan.lower() in ["quit", "q"]:
                print("Exiting the program.")
                sys.exit()
            else:
                print("Feedback noted. Revising the graph...")
                plan_feedback = feedback_plan
                self.memory_manager.results["plan_feedback"] = plan_feedback # Or "plan_feedback_iteration_1", etc.
                print("\n--- Printing self.memory_manager.results in verify and regen ---")
                print(json.dumps(self.memory_manager.results, indent=2)) 
                print("--- End ---")
                logging.info(f"Plan feedback stored in results: {plan_feedback}")
                graph = await self.llm_manager._regenerate_graph_with_feedback(
                    goal, self.tool_executor.tools, feedback_plan, plan.get("graph"), llm_type, selected_agents
                )
                plan["graph"] = graph
                if not plan:
                    print("Failed to revise the plan. Please try again or type 'quit' to exit.")
                    continue
                else:
                    plan_id = plan.get("plan_id")
                    if plan_id:
                        self.plan_manager.plan_storage[plan_id] = {"graph": plan.get("graph"), "goal": goal, "iteration": 0}
                    print("Plan regenerated successfully.")
                    # Continue to the next iteration to re-verify the plan
                    continue

    

    def revise_tools(self, selected_tools: List[Dict], goal: str, feedback: str, llm_type: str = None) -> List[Dict]:
        """Revises the list of selected tools based on user feedback."""
        # Pass the llm_type to _revise_tool_selection
        revised_tools_str = self.llm_manager._revise_tool_selection(selected_tools, goal, feedback, self.tool_executor.tools, llm_type)
        # Parse the LLM's response
        revised_tools = []
        print ("_revise_tool_str.......", revised_tools_str)
        if revised_tools_str:
            tool_names = revised_tools_str.split('\n')
            for tool_name in tool_names:
                tool_name = tool_name.strip()
                if tool_name in self.tool_executor.tools:
                    revised_tools.append({
                        "name": tool_name,
                        "description": self.tool_executor.tools[tool_name]["description"]
                    })
        print ("_revise_tool_str.......", revised_tools)
        return revised_tools

    
    

# --- Helper Functions ---

def _get_tool_descriptions(tools: Dict, selected_agents: List[str] = None) -> str:
    """
    Generates a formatted string containing descriptions of the available tools,
    filtered by selected agents.
    """
    descriptions = ""
    for tool_name, tool_data in tools.items():
        print(f"Processing tool: {tool_name}")
        agent_list = tool_data.get("agent") # using get method and return none if key is missing, prevent from errors
        if selected_agents is None:
           print(f"Adding {tool_name} to descriptions (no selected agents)")
           descriptions += f"{tool_name}: {tool_data['description']}\n"
        elif agent_list is None:
           print(f"Adding {tool_name} to descriptions (agent not in tool_data)")
           descriptions += f"{tool_name}: {tool_data['description']}\n"
        elif "All" in agent_list:
            print(f"Adding {tool_name} to descriptions (agent is 'All')")
            descriptions += f"{tool_name}: {tool_data['description']}\n"
        elif any(agent in agent_list for agent in selected_agents):
            print(f"Adding {tool_name} to descriptions (agent in selected_agents)")
            descriptions += f"{tool_name}: {tool_data['description']}\n"
        else:
            print(f"Skipping {tool_name} (agent not in selected_agents)")

    return descriptions

def _topological_sort(graph: Dict) -> List:
    """
    Performs a topological sort of the graph to determine execution order.
    _topological_sort is designed for generating a linear order that respects all the dependencies in the graph. It visits every node irrespective of the start node.
    Eg:
    {
    "start": ["A"],
    "A": {"next": ["B", "C"]},
    "B": {"next": ["D"]},
    "C": {"next": ["E"]},
    "D": {"next": ["end"]},
    "E": {"next": ["end"]},
     "end": {}
    }
    _topological_sort Order: A possible topological order will be A -> B -> C -> D -> E -> end
    """
    visited = set()
    stack = []

    def visit(node):
        if node not in visited:
            visited.add(node)
            
            # Correctly handle nodes with no outgoing edges
            if isinstance(graph[node], dict) and "next" in graph[node]:
                for neighbor in graph[node]["next"]:
                    visit(neighbor)
            elif isinstance(graph[node], list):
                for neighbor in graph[node]:
                    visit(neighbor)

            stack.append(node)  # Append to the end

    for node in graph:
        visit(node)
    return stack[::-1]  # Reverse the stack for correct order



def build_graph(graph_data):
    """Builds a graph data structure from LLM output, handling dependencies
    correctly, and being robust to various LLM output errors.

    Args:
        graph_data: Dict with the graph structure from the LLM.

    Returns:
        A dictionary with the graph structure, with 'next' lists
        representing dependencies.  Returns an empty dictionary if the input
        is fundamentally flawed.
    """

    graph = {}

    # 1. Create all nodes, handling various potential errors.
    for node_name, node_details in graph_data.items():
        if node_name == "start":
            # Handle "start" node *specifically*.  Its value should be a list.
            if isinstance(node_details, list):
                graph["start"] = node_details
            elif isinstance(node_details, dict) and "next" in node_details:  # Handle dict with "next"
                if isinstance(node_details["next"], list):
                    graph["start"] = node_details["next"]
                elif isinstance(node_details["next"], str):
                    graph["start"] = [node_details["next"]]  # Convert string to list
                else:
                    logging.error(f"Invalid 'start' node format (not a list or string): {node_details}")
                    return {}  # Return empty graph on critical error
            elif isinstance(node_details, str):
                graph["start"] = [node_details]  # Convert string to list.
            else:
                logging.error(f"Invalid 'start' node format (not a list or string): {node_details}")
                return {}  # Return empty graph on critical error.

        elif node_name == "end":
            # Handle end node. Should be {}, but tolerate None.
             graph["end"] = node_details if isinstance(node_details,dict) else {}
        else:
            # Handle regular tool nodes.
            if not isinstance(node_details, dict):
                logging.warning(f"Skipping invalid node '{node_name}': not a dictionary.  Value: {node_details}")
                continue  # Skip this node and move to the next

            if "tool" not in node_details or node_details["tool"] is None:
                logging.warning(f"Skipping node '{node_name}': missing or null 'tool' key.")
                continue

            # Safely get values, providing defaults.  This prevents KeyErrors.
            graph[node_name] = {
                "tool": node_details["tool"],  # We *know* "tool" exists because of the check above
                "input": node_details.get("input", {}),  # Use .get() for safety
                "next": [],  # Initialize 'next' as an empty list
            }

    # 2. Populate the 'next' lists, handling strings and None.
    for node_name, node_details in graph_data.items():
        if node_name == "start" or node_name == "end" or not isinstance(node_details, dict):
            continue

        next_nodes = node_details.get("next")  # Get "next" value, could be None, string, or list
        if isinstance(next_nodes, str):
            graph[node_name]["next"] = [next_nodes]  # Convert string to list
        elif isinstance(next_nodes, list):
            graph[node_name]["next"] = next_nodes  # Use list directly
        elif next_nodes is None:
            graph[node_name]["next"] = []  # Handle None (missing "next") as an empty list
        #No need to raise exception, as we are creating next, if next does not exist.
        #else:  # The only case left should be if its not a string or a list
            #logging.error(f"Invalid 'next' value for node '{node_name}': {next_nodes}")
            #return {}  # Could also choose to skip the node instead of returning empty graph

    # 3. Handle "end" node: Make sure all nodes without outgoing edges connect to "end".
    nodes_with_no_outgoing_edges = [
        node_name for node_name, node_details in graph.items()
        if node_name != "start" and node_name != "end" and not node_details["next"]
    ]
    if "end" not in graph:
        graph["end"] = {}

    for node_name in nodes_with_no_outgoing_edges:
        #This line caused the exception.
        graph[node_name]["next"].append("end") # Directly use next, instead of setdefault

    return graph


def generate_guid():
    """Generates a unique GUID."""
    return str(uuid.uuid4())



async def main():
    config = load_config()
    tools_dir = config["tools_directory"]
                      
    # Configure logging
    log_file_path = os.path.abspath("agent_execution.log")
    file_handler = logging.FileHandler(log_file_path)
    file_handler.setLevel(logging.INFO)  # Set the desired level for the file handler

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(name)s - %(funcName)s - %(message)s",
        handlers=[
            logging.StreamHandler(sys.stdout),  # Output to console
            file_handler,
        ]
    )
    logger = logging.getLogger("agent_framework")
    os.makedirs(tools_dir, exist_ok=True)
    sys.path.insert(0, tools_dir)
    agent =Agent(config)
    
    goal = input("Enter the goal: ")
    payload_data = {"goal": goal, "llm_type": 'gemini'}
    payload = RunAgentRequest(**payload_data)
    session_id = generate_guid()
    result = await agent.run_agent(payload, session_id)
    print(result)  
    
    # Run the async function using asyncio
if __name__ == "__main__":
    asyncio.run(main())
